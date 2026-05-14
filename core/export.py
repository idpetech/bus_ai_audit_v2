"""
Export functionality for BA Assistant
PDF and Word document generation
"""

import io
import logging
import re
from datetime import datetime
from typing import TYPE_CHECKING

import markdown
from bs4 import BeautifulSoup
from fpdf import FPDF
from docx import Document

if TYPE_CHECKING:
    from .models import PipelineResults

logger = logging.getLogger(__name__)


class PDFGenerator:
    def __init__(self):
        pass
    
    def _clean_text(self, text: str) -> str:
        """Remove markdown and clean text for PDF"""
        # Convert markdown to HTML then extract text
        html = markdown.markdown(text)
        soup = BeautifulSoup(html, 'html.parser')
        cleaned = soup.get_text()
        # Remove extra whitespace
        cleaned = re.sub(r'\n\s*\n', '\n\n', cleaned)
        # Replace problematic characters using explicit Unicode code points
        cleaned = cleaned.replace('\u201c', '"').replace('\u201d', '"')
        cleaned = cleaned.replace('\u2018', "'").replace('\u2019', "'")
        cleaned = cleaned.replace('–', '-').replace('—', '-')
        cleaned = cleaned.replace('…', '...')
        # Remove any remaining non-ASCII characters
        cleaned = ''.join(char if ord(char) < 128 else '?' for char in cleaned)
        return cleaned.strip()
    
    def _add_content_to_pdf(self, pdf: FPDF, title: str, content: str, company_name: str = None):
        """Add content section to PDF"""
        # Header
        pdf.set_font('Times', 'B', 16)
        pdf.cell(0, 10, f'IDPETECH - BA Assistant: {title}', ln=True, align='C')
        pdf.set_font('Times', '', 10)
        pdf.cell(0, 10, f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}', ln=True, align='C')
        pdf.ln(5)
        
        # Company name if provided
        if company_name:
            pdf.set_font('Times', 'B', 14)
            pdf.cell(0, 10, f'Company: {company_name}', ln=True)
            pdf.ln(5)
        
        # Content
        pdf.set_font('Times', '', 10)
        cleaned_content = self._clean_text(content)
        
        # Split long text into lines
        lines = cleaned_content.split('\n')
        for line in lines:
            if len(line) > 80:
                # Wrap long lines
                words = line.split(' ')
                current_line = ""
                for word in words:
                    if len(current_line + word) < 80:
                        current_line += word + " "
                    else:
                        if current_line:
                            pdf.cell(0, 5, current_line.strip(), ln=True)
                        current_line = word + " "
                if current_line:
                    pdf.cell(0, 5, current_line.strip(), ln=True)
            else:
                pdf.cell(0, 5, line, ln=True)
    
    def generate_section_pdf(self, title: str, content: str, company_name: str = None) -> bytes:
        """Generate PDF for individual section"""
        pdf = FPDF()
        pdf.add_page()
        self._add_content_to_pdf(pdf, title, content, company_name)
        return bytes(pdf.output())
    
    def generate_pdf(self, results: 'PipelineResults', company_name: str) -> bytes:
        """Generate complete PDF report"""
        pdf = FPDF()
        pdf.add_page()
        
        # Header for complete report
        pdf.set_font('Times', 'B', 16)
        pdf.cell(0, 10, 'IDPETECH - BA Assistant: Complete AI Readiness Report', ln=True, align='C')
        pdf.set_font('Times', '', 10)
        pdf.cell(0, 10, f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}', ln=True, align='C')
        pdf.ln(5)
        
        # Company name
        pdf.set_font('Times', 'B', 14)
        pdf.cell(0, 10, f'Company: {company_name}', ln=True)
        pdf.ln(5)
        
        sections = [
            ("First Engagement Hook", results.hook),
            ("Diagnostic Insight", results.diagnosis),
            ("AI Readiness Audit", results.audit),
            ("Conversation Close", results.close)
        ]
        
        for title, content in sections:
            pdf.set_font('Times', 'B', 12)
            pdf.cell(0, 10, title, ln=True)
            pdf.ln(2)
            
            pdf.set_font('Times', '', 10)
            cleaned_content = self._clean_text(content)
            
            # Split long text into lines
            lines = cleaned_content.split('\n')
            for line in lines:
                if len(line) > 80:
                    # Wrap long lines
                    words = line.split(' ')
                    current_line = ""
                    for word in words:
                        if len(current_line + word) < 80:
                            current_line += word + " "
                        else:
                            if current_line:
                                pdf.cell(0, 5, current_line.strip(), ln=True)
                            current_line = word + " "
                    if current_line:
                        pdf.cell(0, 5, current_line.strip(), ln=True)
                else:
                    pdf.cell(0, 5, line, ln=True)
            pdf.ln(5)
        
        return bytes(pdf.output())


class WordGenerator:
    def __init__(self):
        pass
    
    def _clean_text_for_word(self, text: str) -> str:
        """Clean and prepare text for Word document"""
        # Convert markdown to plain text while preserving structure
        html = markdown.markdown(text)
        soup = BeautifulSoup(html, 'html.parser')
        cleaned = soup.get_text()
        
        # Clean up spacing
        cleaned = re.sub(r'\n\s*\n', '\n\n', cleaned)
        return cleaned.strip()
    
    def _add_markdown_content(self, doc: Document, content: str):
        """Add markdown content to Word document with proper formatting"""
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
                
            # Handle headers
            if line.startswith('# '):
                heading = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                heading = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                heading = doc.add_heading(line[4:], level=3)
            elif line.startswith('**') and line.endswith('**'):
                # Bold text
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
            elif line.startswith('- '):
                # Bullet point
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                # Regular paragraph
                doc.add_paragraph(line)
            
            i += 1
    
    def generate_section_word(self, title: str, content: str, company_name: str = None) -> bytes:
        """Generate Word document for individual section"""
        doc = Document()
        
        # Add header
        header_p = doc.add_heading(f'IDPETECH - BA Assistant: {title}', level=1)
        
        # Add metadata
        meta_p = doc.add_paragraph()
        meta_p.add_run(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}\n').italic = True
        if company_name:
            meta_p.add_run(f'Company: {company_name}').italic = True
        
        doc.add_paragraph()  # Space
        
        # Add content
        self._add_markdown_content(doc, content)
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def generate_word(self, results: 'PipelineResults', company_name: str) -> bytes:
        """Generate complete Word document report"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('IDPETECH - BA Assistant: Complete AI Readiness Report', level=1)
        
        # Add metadata
        meta_p = doc.add_paragraph()
        meta_p.add_run(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}\n').italic = True
        meta_p.add_run(f'Company: {company_name}').italic = True
        
        doc.add_page_break()
        
        sections = [
            ("First Engagement Hook", results.hook),
            ("Diagnostic Insight", results.diagnosis),
            ("AI Readiness Audit", results.audit),
            ("Conversation Close", results.close)
        ]
        
        for title, content in sections:
            doc.add_heading(title, level=2)
            doc.add_paragraph()  # Space
            self._add_markdown_content(doc, content)
            doc.add_page_break()
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()