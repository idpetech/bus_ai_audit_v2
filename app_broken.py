import streamlit as st
import openai
import json
import hashlib
import logging
import time
import os
from datetime import datetime
from typing import Dict, Any, Optional, Tuple, List
from dataclasses import dataclass, field
from urllib.parse import urlparse
import requests
import html2text
from fpdf import FPDF
import markdown
from bs4 import BeautifulSoup
import re
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
import io
import tiktoken

# Import extracted core components
from core.models import CompanyInputs, PipelineResults
from core.utils import _is_url, sieve_context
from core.database import DatabaseManager
from core.scraping import FirecrawlManager, scrape_website, scrape_page
from core.pipeline import BAAssistant
from core.export import PDFGenerator, WordGenerator

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)













    def __init__(self, api_key: str, prompts: Dict[str, str] = None):
        self.client = openai.OpenAI(api_key=api_key)
        self.cache = {}
        self.prompts = prompts or self._load_prompts()
    
    def _get_default_prompts(self) -> Dict[str, str]:
        """Default system prompts for each stage"""
        return {
            "extract_signals": """You are a signal extraction engine. Extract ONLY structured data from company inputs.

Return valid JSON with these fields:
{
  "company_name": "",
  "industry": "",
  "stage": "",
  "tech_stack": [],
  "role_being_hired": "",
  "role_seniority": "",
  "specific_skills_required": [],
  "business_model_signals": [],
  "scale_indicators": [],
  "ai_mentions": [],
  "technical_complexity_signals": [],
  "architecture_keywords": [],
  "data_flow_indicators": [],
  "scaling_evidence": []
}

NO interpretation. NO analysis. ONLY data extraction.""",
            
            "diagnose": """You are a Senior Principal Architect with 35+ years building distributed systems. You've seen every architectural failure mode and can spot BS from orbit.

CRITICAL: Use triangulated data to expose the gap between marketing narratives and engineering reality.

REQUIRED OUTPUT STRUCTURE:

**What They Think They're Building:**
[Extract from company narrative - their vision, marketing claims, stated technical goals]

**What They're Actually Building:**
[Based on external signals - actual tech stack, hiring patterns, engineering challenges, reviews]

**The Architectural Gap:**
[Specific technical contradictions between claimed capabilities and engineering reality]

**System Failure Modes:**
- Data flow bottlenecks that will break at scale
- Technical debt blocking AI implementation pathways
- Operational complexity gaps they're blind to
- Team capability mismatches vs. stated technical ambitions

**Breaking Points at 10x Scale:**
[Specific architectural chokepoints that will fail based on external signals]

**AI Implementation Reality:**
- AI-Washed: Pure marketing play, traditional CRUD underneath
- AI-Assisted: Bolting AI onto existing architectures
- AI-Native: Core business logic requires AI to function
- Non-AI: Traditional software, no meaningful AI dependency

**Technical Classification Evidence:**
[Justify with concrete evidence from external signals vs. company claims]

Be brutally honest. No sugar-coating. Focus on what breaks systems.""",
            
            "generate_hook": """You are a Senior Principal Architect reaching out peer-to-peer. No introduction. No fluff. Start with a direct architectural observation.

REQUIREMENTS:
- 2-4 sentences maximum
- Lead with specific technical contradiction (hiring vs. claims vs. stack)
- Frame as architectural curiosity, not accusation
- Reference concrete external signals (not generic observations)
- End with direct technical question

TONE: Senior technical peer. Skeptical but respectful.

AVOID:
- Any introduction ("I noticed", "Hi", "Hope you're well")
- Buzzwords (scale, optimize, leverage, unlock, transform)
- Generic AI statements
- Sales language or multiple topics

FORMULA:
[Technical observation] + [Specific contradiction] + [Direct question]

Example: "Your team is hiring Node.js engineers while positioning as an AI-first company, but I don't see ML infrastructure in the stack. What's the actual data flow architecture you're building?"

Start immediately with the technical tension. No preamble.
""",
            
            "generate_audit": """You are a Senior Principal Architect conducting a brutal technical assessment. Focus on failure modes and architectural realities.

# AI Readiness Audit

## Executive Summary
- Current AI Classification: [AI-Washed/AI-Assisted/AI-Native/Non-AI]
- Reality Gap Score: X/10 (claims vs. actual technical capability)
- Primary Constraint: [The one architectural bottleneck that will break first]

## System Reality
What they actually built vs what they claim to be building.

## Architectural Constraint Analysis
**Data Architecture:** [Specific data flow bottlenecks and pipeline limitations]
**Infrastructure Reality:** [Actual processing constraints vs. AI requirements]  
**Team Technical Debt:** [Skills gaps that will cause implementation failures]
**Legacy System Constraints:** [Technical debt blocking AI integration pathways]
**Operational Maturity Gaps:** [Process/tooling failures at scale]

## Failure Mode Predictions
Top 3 ways their AI implementation will break, with specific technical triggers.

## Contrarian Technical Assessment
**What consultants are telling them:** [Standard advice they're receiving]
**Architectural reality:** [Uncomfortable technical truths they're avoiding]
**The technical bet:** [What must be architecturally true for success]

## Mechanism-Based Interventions
Concrete technical changes, not strategy:
- [Specific architectural modifications with failure prevention]
- [Data flow optimizations with measurable outcomes]
- [Infrastructure decisions that prevent known failure modes]

No buzzwords. No "best practices." Focus on what breaks distributed systems.""",
            
            "generate_close": """You are a Senior Architect delivering a final technical observation. Cut through the noise.

Requirements:
- 2-3 sentences maximum
- Surface one core architectural contradiction
- No questions, no offers, no next steps
- End with direct statement: what they're building vs. what they think they're building

Avoid: consulting language, encouragement, solutions, positive spin

Tone: Peer-level technical honesty. No sugar-coating.

Example: "Your AI roadmap assumes clean data flows, but your hiring suggests you're still fighting legacy ETL pipelines. You're building a data cleanup operation, not an AI product."

Direct. Final. Architectural truth."""
        }
    
    def _load_prompts(self) -> Dict[str, str]:
        """Load prompts from custom file if exists, otherwise use defaults"""
        custom_prompts_file = "custom_prompts.json"
        if os.path.exists(custom_prompts_file):
            try:
                with open(custom_prompts_file, 'r') as f:
                    return json.load(f)
            except Exception as e:
                logger.warning(f"Could not load custom prompts: {e}")
        return self._get_default_prompts()
    
    def save_custom_prompts(self, prompts: Dict[str, str]) -> bool:
        """Save custom prompts as new defaults"""
        try:
            with open("custom_prompts.json", 'w') as f:
                json.dump(prompts, f, indent=2)
            return True
        except Exception as e:
            logger.error(f"Could not save custom prompts: {e}")
            return False
    
    def reset_to_factory_defaults(self) -> Dict[str, str]:
        """Reset to factory defaults and remove custom file"""
        try:
            if os.path.exists("custom_prompts.json"):
                os.remove("custom_prompts.json")
        except Exception as e:
            logger.warning(f"Could not remove custom prompts file: {e}")
        return self._get_default_prompts()
    
    def _make_llm_call(self, system_prompt: str, user_prompt: str, model: str = "gpt-4o", max_retries: int = 3, stage: str = "unknown") -> str:
        """Reusable LLM wrapper with retry handling and TPM management"""
        
        # SIGNAL SIEVE: Apply context truncation to stay under TPM limits
        from core.utils import MAX_CONTEXT_TOKENS
        sieved_prompt = sieve_context(user_prompt, MAX_CONTEXT_TOKENS)
        
        # TEMPERATURE CONTROL: Technical precision for critical stages
        temp = 0.4 if stage in ["diagnose", "generate_hook"] else 0.7
        
        for attempt in range(max_retries):
            try:
                logger.info(f"🤖 {stage.upper()}: Making LLM call (temp={temp})")
                response = self.client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": sieved_prompt}
                    ],
                    temperature=temp
                )
                return response.choices[0].message.content
            except Exception as e:
                logger.error(f"LLM call attempt {attempt + 1} failed: {e}")
                if attempt == max_retries - 1:
                    raise e
                time.sleep(2 ** attempt)
    
    def _get_cache_key(self, inputs: CompanyInputs) -> str:
        """Generate cache key from inputs"""
        combined = f"{inputs.target_url}{inputs.job_posting or ''}"
        return hashlib.md5(combined.encode()).hexdigest()

    def _extract_json(self, text: str) -> Dict[str, Any]:
        """Parse JSON from LLM response, stripping markdown code fences if present."""
        stripped = re.sub(r"^```(?:json)?\s*|\s*```$", "", text.strip(), flags=re.DOTALL)
        return json.loads(stripped)

    def extract_signals(self, inputs: CompanyInputs) -> Dict[str, Any]:
        """Stage 1: Extract structured signals without interpretation"""
        system_prompt = self.prompts["extract_signals"]

        user_prompt = f"""Extract signals from triangulated company data:

Target URL: {inputs.target_url}
Company: {inputs.company_name or 'Unknown'}

{inputs.combined_context}"""

        response = self._make_llm_call(system_prompt, user_prompt, stage="extract_signals")
        logger.info(f"Signals extracted: {response[:200]}...")

        try:
            return self._extract_json(response)
        except json.JSONDecodeError:
            return {"error": "Failed to parse signals", "raw_response": response}
    
    def diagnose(self, signals: Dict[str, Any], inputs: CompanyInputs) -> str:
        """Stage 2: Agentic reality check using triangulated data"""
        system_prompt = self.prompts["diagnose"]

        user_prompt = f"""Perform reality check on company using triangulated intelligence:

Company: {inputs.company_name or 'Unknown'}
URL: {inputs.target_url}

EXTRACTED SIGNALS: {json.dumps(signals, indent=2)}

TRIANGULATED DATA:
{inputs.combined_context}

FOCUS: Use external signals to challenge company self-perception. Expose contradictions between what they claim vs what external sources reveal about their technical reality."""

        response = self._make_llm_call(system_prompt, user_prompt, max_retries=3, stage="diagnose")
        logger.info(f"Agentic diagnosis completed: {response[:200]}...")
        return response
    
    def generate_hook(self, signals: Dict[str, Any], diagnosis: str) -> str:
        """Stage 3: Generate founder-facing outbound hook"""
        system_prompt = self.prompts["generate_hook"]

        user_prompt = f"""Create hook message based on:

Signals: {json.dumps(signals, indent=2)}
Diagnosis: {diagnosis}"""

        response = self._make_llm_call(system_prompt, user_prompt, stage="generate_hook")
        logger.info("Hook generated")
        return response
    
    def generate_audit(self, signals: Dict[str, Any], diagnosis: str) -> str:
        """Stage 4: Generate structured audit report"""
        system_prompt = self.prompts["generate_audit"]

        user_prompt = f"""Generate audit for:

Signals: {json.dumps(signals, indent=2)}
Diagnosis: {diagnosis}"""

        response = self._make_llm_call(system_prompt, user_prompt, stage="generate_audit")
        logger.info("Audit generated")
        return response
    
    def generate_close(self, signals: Dict[str, Any], audit: str) -> str:
        """Stage 5: Generate conversation close with soft CTA"""
        system_prompt = self.prompts["generate_close"]

        user_prompt = f"""Create close based on:

Company Signals: {json.dumps(signals, indent=2)}
Audit Summary: {audit[:500]}..."""

        response = self._make_llm_call(system_prompt, user_prompt, stage="generate_close")
        logger.info("Close generated")
        return response
    
    def run_full_pipeline(self, inputs: CompanyInputs) -> PipelineResults:
        """Execute complete 5-stage pipeline"""
        cache_key = self._get_cache_key(inputs)
        
        if cache_key in self.cache:
            logger.info("Returning cached results")
            return self.cache[cache_key]
        
        # Stage 1: Extract signals
        logger.info(f"🔄 Stage 1/5: Extracting structured signals...")
        signals = self.extract_signals(inputs)
        
        # WAIT & PULSE: TPM rate limiting between extract and diagnose
        logger.info(f"⏳ Waiting {_PIPELINE_DELAY}s for TPM bucket refill...")
        time.sleep(_PIPELINE_DELAY)
        
        # Stage 2: Diagnose
        logger.info(f"🔄 Stage 2/5: Running agentic diagnosis...")
        diagnosis = self.diagnose(signals, inputs)
        
        # Stage 3: Generate hook
        logger.info(f"🔄 Stage 3/5: Generating founder hook...")
        hook = self.generate_hook(signals, diagnosis)
        
        # Stage 4: Generate audit
        logger.info(f"🔄 Stage 4/5: Building audit report...")
        audit = self.generate_audit(signals, diagnosis)
        
        # Stage 5: Generate close
        logger.info(f"🔄 Stage 5/5: Finalizing conversation close...")
        close = self.generate_close(signals, audit)
        
        results = PipelineResults(
            signals=signals,
            diagnosis=diagnosis,
            hook=hook,
            audit=audit,
            close=close
        )
        
        # Cache results
        self.cache[cache_key] = results
        logger.info("Pipeline completed and cached")
        
        return results
    
    def run_triangulation(self, url: str, job_posting: Optional[str] = None, 
                         firecrawl_manager: FirecrawlManager = None, 
                         db_manager: DatabaseManager = None) -> Tuple[CompanyInputs, PipelineResults]:
        """
        Agentic Triangulation Loop: Automated reality check engine
        
        Step 1 (Scrape): Extract company's self-perception narrative
        Step 2 (Search): Hunt for external signals about technical reality  
        Step 3 (Cross-Reference): Feed both into diagnosis for contradiction analysis
        
        Returns: (CompanyInputs with triangulated data, PipelineResults)
        """
        logger.info(f"🤖 Starting agentic triangulation for {url}")
        
        # Step 1: Scrape company narrative (self-perception)
        logger.info("⚗️ Step 1: Distilling company self-perception narrative...")
        scrape_success, scraped_content, company_name = firecrawl_manager.scrape_company_narrative(url)
        
        if not scrape_success:
            logger.error(f"Failed to scrape {url}: {scraped_content}")
            raise Exception(f"Scraping failed: {scraped_content}")
        
        # Step 2: Search for external signals (technical reality)
        logger.info("🔍 Step 2: Hunting for external technical contradiction signals...")
        search_success, external_signals = firecrawl_manager.search_external_signals(company_name, url)
        
        if not search_success:
            logger.warning(f"No external signals found: {external_signals}")
            external_signals = "No external signals found. Analysis will be based solely on company narrative."
        
        # Step 3: Cross-reference both sources for reality check
        logger.info("⚖️ Step 3: Verifying tech-stack contradictions vs. company claims...")
        
        # Create enriched inputs with triangulated data
        inputs = CompanyInputs(
            target_url=url,
            job_posting=job_posting,
            scraped_content=scraped_content,
            external_signals=external_signals,
            company_name=company_name
        )
        
        # Run the analysis pipeline with triangulated data
        results = self.run_full_pipeline(inputs)
        
        # Persist to database if manager provided
        if db_manager:
            db_manager.upsert_analysis(url, inputs, results)
            logger.info(f"💾 Analysis persisted to database for {url}")
        
        logger.info(f"🎯 Triangulation complete for {company_name}")
        return inputs, results

class PDFGenerator:
    def __init__(self):
        pass
    
    def _clean_text(self, text: str) -> str:
        """Remove markdown and clean text for PDF"""
        # Convert markdown to HTML then extract text
        html = markdown.markdown(text)
        soup = BeautifulSoup(html, 'html.parser')
        cleaned = soup.get_text()
        # Remove extra whitespace
        cleaned = re.sub(r'\n\s*\n', '\n\n', cleaned)
        # Replace problematic characters using explicit Unicode code points
        cleaned = cleaned.replace('\u201c', '"').replace('\u201d', '"')
        cleaned = cleaned.replace('\u2018', "'").replace('\u2019', "'")
        cleaned = cleaned.replace('–', '-').replace('—', '-')
        cleaned = cleaned.replace('…', '...')
        # Remove any remaining non-ASCII characters
        cleaned = ''.join(char if ord(char) < 128 else '?' for char in cleaned)
        return cleaned.strip()
    
    def _add_content_to_pdf(self, pdf: FPDF, title: str, content: str, company_name: str = None):
        """Add content section to PDF"""
        # Header
        pdf.set_font('Times', 'B', 16)
        pdf.cell(0, 10, f'IDPETECH - BA Assistant: {title}', ln=True, align='C')
        pdf.set_font('Times', '', 10)
        pdf.cell(0, 10, f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}', ln=True, align='C')
        pdf.ln(5)
        
        # Company name if provided
        if company_name:
            pdf.set_font('Times', 'B', 14)
            pdf.cell(0, 10, f'Company: {company_name}', ln=True)
            pdf.ln(5)
        
        # Content
        pdf.set_font('Times', '', 10)
        cleaned_content = self._clean_text(content)
        
        # Split long text into lines
        lines = cleaned_content.split('\n')
        for line in lines:
            if len(line) > 80:
                # Wrap long lines
                words = line.split(' ')
                current_line = ""
                for word in words:
                    if len(current_line + word) < 80:
                        current_line += word + " "
                    else:
                        if current_line:
                            pdf.cell(0, 5, current_line.strip(), ln=True)
                        current_line = word + " "
                if current_line:
                    pdf.cell(0, 5, current_line.strip(), ln=True)
            else:
                pdf.cell(0, 5, line, ln=True)
    
    def generate_section_pdf(self, title: str, content: str, company_name: str = None) -> bytes:
        """Generate PDF for individual section"""
        pdf = FPDF()
        pdf.add_page()
        self._add_content_to_pdf(pdf, title, content, company_name)
        return bytes(pdf.output())
    
    def generate_pdf(self, results: PipelineResults, company_name: str) -> bytes:
        """Generate complete PDF report"""
        pdf = FPDF()
        pdf.add_page()
        
        # Header for complete report
        pdf.set_font('Times', 'B', 16)
        pdf.cell(0, 10, 'IDPETECH - BA Assistant: Complete AI Readiness Report', ln=True, align='C')
        pdf.set_font('Times', '', 10)
        pdf.cell(0, 10, f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}', ln=True, align='C')
        pdf.ln(5)
        
        # Company name
        pdf.set_font('Times', 'B', 14)
        pdf.cell(0, 10, f'Company: {company_name}', ln=True)
        pdf.ln(5)
        
        sections = [
            ("First Engagement Hook", results.hook),
            ("Diagnostic Insight", results.diagnosis),
            ("AI Readiness Audit", results.audit),
            ("Conversation Close", results.close)
        ]
        
        for title, content in sections:
            pdf.set_font('Times', 'B', 12)
            pdf.cell(0, 10, title, ln=True)
            pdf.ln(2)
            
            pdf.set_font('Times', '', 10)
            cleaned_content = self._clean_text(content)
            
            # Split long text into lines
            lines = cleaned_content.split('\n')
            for line in lines:
                if len(line) > 80:
                    # Wrap long lines
                    words = line.split(' ')
                    current_line = ""
                    for word in words:
                        if len(current_line + word) < 80:
                            current_line += word + " "
                        else:
                            if current_line:
                                pdf.cell(0, 5, current_line.strip(), ln=True)
                            current_line = word + " "
                    if current_line:
                        pdf.cell(0, 5, current_line.strip(), ln=True)
                else:
                    pdf.cell(0, 5, line, ln=True)
            pdf.ln(5)
        
        return bytes(pdf.output())

class WordGenerator:
    def __init__(self):
        pass
    
    def _clean_text_for_word(self, text: str) -> str:
        """Clean and prepare text for Word document"""
        # Convert markdown to plain text while preserving structure
        html = markdown.markdown(text)
        soup = BeautifulSoup(html, 'html.parser')
        cleaned = soup.get_text()
        
        # Clean up spacing
        cleaned = re.sub(r'\n\s*\n', '\n\n', cleaned)
        return cleaned.strip()
    
    def _add_markdown_content(self, doc: Document, content: str):
        """Add markdown content to Word document with proper formatting"""
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
                
            # Handle headers
            if line.startswith('# '):
                heading = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                heading = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                heading = doc.add_heading(line[4:], level=3)
            elif line.startswith('**') and line.endswith('**'):
                # Bold text
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
            elif line.startswith('- '):
                # Bullet point
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                # Regular paragraph
                doc.add_paragraph(line)
            
            i += 1
    
    def generate_section_word(self, title: str, content: str, company_name: str = None) -> bytes:
        """Generate Word document for individual section"""
        doc = Document()
        
        # Add header
        header_p = doc.add_heading(f'IDPETECH - BA Assistant: {title}', level=1)
        
        # Add metadata
        meta_p = doc.add_paragraph()
        meta_p.add_run(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}\n').italic = True
        if company_name:
            meta_p.add_run(f'Company: {company_name}').italic = True
        
        doc.add_paragraph()  # Space
        
        # Add content
        self._add_markdown_content(doc, content)
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def generate_word(self, results: PipelineResults, company_name: str) -> bytes:
        """Generate complete Word document report"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('IDPETECH - BA Assistant: Complete AI Readiness Report', level=1)
        
        # Add metadata
        meta_p = doc.add_paragraph()
        meta_p.add_run(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}\n').italic = True
        meta_p.add_run(f'Company: {company_name}').italic = True
        
        doc.add_page_break()
        
        sections = [
            ("First Engagement Hook", results.hook),
            ("Diagnostic Insight", results.diagnosis),
            ("AI Readiness Audit", results.audit),
            ("Conversation Close", results.close)
        ]
        
        for title, content in sections:
            doc.add_heading(title, level=2)
            doc.add_paragraph()  # Space
            self._add_markdown_content(doc, content)
            doc.add_page_break()
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

def main():
    st.set_page_config(
        page_title="IDPETECH · Architectural Reality Check Engine",
        page_icon="⚗️",
        layout="wide"
    )
    
    st.title("⚗️ IDPETECH · Architectural Reality Check Engine")
    st.subheader("Senior Principal Architect · AI Claims vs Technical Reality Assessment")
    
    # Initialize session state
    if 'results' not in st.session_state:
        st.session_state.results = None
    if 'inputs' not in st.session_state:
        st.session_state.inputs = None
    if 'prompts' not in st.session_state:
        ba_temp = BAAssistant("temp")
        st.session_state.prompts = ba_temp._load_prompts()
    
    # Initialize managers
    openai_api_key = st.secrets.get("OPENAI_API_KEY") or st.sidebar.text_input("OpenAI API Key", type="password")
    firecrawl_api_key = st.secrets.get("FIRECRAWL_API_KEY") or st.sidebar.text_input("Firecrawl API Key", type="password")
    
    if not openai_api_key:
        st.warning("Please provide OpenAI API key to continue.")
        return
    if not firecrawl_api_key:
        st.warning("Please provide Firecrawl API key to continue.")
        return
    
    # Initialize managers
    if 'ba_assistant' not in st.session_state:
        st.session_state.ba_assistant = BAAssistant(openai_api_key, st.session_state.prompts)
    if 'firecrawl_manager' not in st.session_state:
        st.session_state.firecrawl_manager = FirecrawlManager(firecrawl_api_key)
    if 'db_manager' not in st.session_state:
        st.session_state.db_manager = DatabaseManager()
    
    # Layout: Left panel for inputs, Right panel for outputs
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.header("⚗️ Architectural Assessment")
        
        # Company database browser
        with st.expander("📚 Company Database", expanded=False):
            companies = st.session_state.db_manager.list_companies()
            if companies:
                st.write(f"**{len(companies)} companies analyzed:**")
                for url, name, updated in companies[:10]:  # Show latest 10
                    col_name, col_actions, col_date = st.columns([2.5, 2.5, 1])
                    with col_name:
                        st.write(f"**{name}**")
                    with col_actions:
                        # Create three loading buttons in a single row
                        btn_col1, btn_col2, btn_col3, btn_col4 = st.columns([1, 1, 1, 0.8])
                        
                        with btn_col1:
                            if st.button("📄", key=f"context_{hashlib.md5(url.encode()).hexdigest()[:8]}", help="Load context only (for reprocessing)"):
                                # Load only context for reprocessing
                                context_data = st.session_state.db_manager.get_context_only(url)
                                if context_data:
                                    st.session_state.inputs, _ = context_data
                                    st.session_state.results = None  # Clear previous results
                                    st.success(f"Loaded context for {name} - ready for reprocessing")
                                    st.rerun()
                        
                        with btn_col2:
                            if st.button("📊", key=f"results_{hashlib.md5(url.encode()).hexdigest()[:8]}", help="Load results only (view analysis)"):
                                # Load only results for viewing
                                cached_data = st.session_state.db_manager.get_analysis(url)
                                if cached_data:
                                    _, st.session_state.results, _ = cached_data
                                    st.session_state.inputs = None  # Clear inputs
                                    st.success(f"Loaded results for {name}")
                                    st.rerun()
                        
                        with btn_col3:
                            if st.button("📂", key=f"full_{hashlib.md5(url.encode()).hexdigest()[:8]}", help="Load complete analysis (context + results)"):
                                # Load complete analysis
                                cached_data = st.session_state.db_manager.get_analysis(url)
                                if cached_data:
                                    st.session_state.inputs, st.session_state.results, _ = cached_data
                                    st.success(f"Loaded complete analysis for {name}")
                                    st.rerun()
                        
                        with btn_col4:
                            if st.button("🗑️", key=f"delete_{hashlib.md5(url.encode()).hexdigest()[:8]}", help="Delete company record"):
                                if st.session_state.db_manager.delete_company(url):
                                    st.success(f"Deleted {name}")
                                    st.rerun()
                                else:
                                    st.error(f"Failed to delete {name}")
                    
                    with col_date:
                        st.caption(updated.strftime("%m/%d"))
                        
                # Add comprehensive legend for buttons
                st.caption("**Load Options:** 📄 Context only (reprocess) • 📊 Results only (view) • 📂 Complete analysis • 🗑️ Delete record")
            else:
                st.info("No companies analyzed yet.")
        
        st.divider()
        
        # Primary inputs
        target_url = st.text_input(
            "🌐 Target Company URL",
            placeholder="https://company.com",
            help="Primary company website for automated analysis"
        )
        
        job_posting = st.text_area(
            "💼 Job Posting (Optional)",
            placeholder="Paste job posting text for additional context...",
            height=120,
            help="Optional job posting to enhance technical signal detection"
        )
        
        # Context editing section (shows when context is loaded for reprocessing)
        if 'inputs' in st.session_state and st.session_state.inputs and not st.session_state.get('results'):
            st.divider()
            with st.expander("✏️ Edit Loaded Context", expanded=True):
                st.info("📝 Context loaded from database. You can edit and add additional context before reprocessing.")
                
                # Show/edit scraped content
                if st.session_state.inputs.scraped_content:
                    st.subheader("Company Self-Perception")
                    edited_scraped = st.text_area(
                        "Company narrative (scraped content):",
                        value=st.session_state.inputs.scraped_content,
                        height=150,
                        key="edit_scraped"
                    )
                    st.session_state.inputs.scraped_content = edited_scraped
                
                # Show/edit external signals
                if st.session_state.inputs.external_signals:
                    st.subheader("External Technical Signals")
                    edited_signals = st.text_area(
                        "External signals:",
                        value=st.session_state.inputs.external_signals,
                        height=150,
                        key="edit_signals"
                    )
                    st.session_state.inputs.external_signals = edited_signals
                
                # Additional context input
                st.subheader("Additional Context")
                additional_context = st.text_area(
                    "Add any additional context or insights:",
                    placeholder="Add new technical insights, competitor analysis, or other relevant context...",
                    height=100,
                    key="additional_context"
                )
                
                # Append additional context to job posting if provided
                if additional_context and additional_context.strip():
                    if st.session_state.inputs.job_posting:
                        st.session_state.inputs.job_posting = f"{st.session_state.inputs.job_posting}\n\nAdditional Context:\n{additional_context}"
                    else:
                        st.session_state.inputs.job_posting = f"Additional Context:\n{additional_context}"
                
                # Override the primary inputs when editing
                if st.session_state.inputs.target_url:
                    target_url = st.session_state.inputs.target_url
                if st.session_state.inputs.job_posting:
                    job_posting = st.session_state.inputs.job_posting
        
        # Check for cached analysis
        cached_analysis = None
        if target_url:
            cached_analysis = st.session_state.db_manager.get_analysis(target_url)
        
        # Main action buttons
        run_analysis = False
        
        # Handle different loading modes
        if 'inputs' in st.session_state and st.session_state.inputs and not st.session_state.get('results'):
            # Context-only mode (ready for reprocessing)
            st.info("🔄 **Context Loaded for Reprocessing** - Ready to run analysis on existing context")
            if st.button("⚗️ Reprocess with Context", type="primary", help="Run analysis on loaded context (uses API credits)"):
                run_analysis = True
        elif 'results' in st.session_state and st.session_state.results and not st.session_state.get('inputs'):
            # Results-only mode (view only)
            st.info("📊 **Results Loaded for Viewing** - Analysis results displayed below")
            st.caption("💡 To reprocess this analysis, use the 📄 Context button to load context for editing")
        elif cached_analysis:
            col_run, col_refresh = st.columns(2)
            with col_run:
                if st.button("⚗️ Run Fresh Assessment", type="primary", help="Run new technical analysis (uses API credits)"):
                    run_analysis = True
            with col_refresh:
                if st.button("📂 Load Cached", help="Load existing analysis from database"):
                    st.session_state.inputs, st.session_state.results, last_updated = cached_analysis
                    st.success(f"Loaded analysis from {last_updated.strftime('%Y-%m-%d %H:%M')}")
                    st.rerun()
        else:
            if st.button("⚗️ Start Architecture Analysis", type="primary", help="Distill + Hunt + Cross-Reference"):
                run_analysis = True
        
        if run_analysis:
            # Check if we're reprocessing existing context or running fresh analysis
            if 'inputs' in st.session_state and st.session_state.inputs and not st.session_state.get('results'):
                # Context reprocessing mode
                if not target_url:
                    st.error("Please provide a target company URL.")
                elif not _is_url(target_url):
                    st.error("Please provide a valid URL (must start with http:// or https://)")
                else:
                    # Status tracking for context reprocessing
                    status_container = st.status("🔄 Reprocessing Context with Senior Architect Analysis...", expanded=True)
                    
                    try:
                        with status_container:
                            st.write("📝 Using loaded context from database...")
                            st.write("⚗️ Running pipeline on existing context...")
                            st.write("🎯 Generating fresh insights...")
                            
                            # Run pipeline on existing context
                            results = st.session_state.ba_assistant.run_full_pipeline(st.session_state.inputs)
                            st.session_state.results = results
                            
                            # Save updated analysis to database
                            st.session_state.db_manager.upsert_analysis(
                                st.session_state.inputs.target_url, 
                                st.session_state.inputs, 
                                results
                            )
                    except Exception as e:
                        status_container.update(
                            label="❌ Context Reprocessing Failed", 
                            state="error", 
                            expanded=True
                        )
                        st.error(f"Context reprocessing failed: {str(e)}")
                        logger.error(f"Context reprocessing failed: {e}")
                        
            else:
                # Fresh analysis mode
                if not target_url:
                    st.error("Please provide a target company URL.")
                elif not _is_url(target_url):
                    st.error("Please provide a valid URL (must start with http:// or https://)")
                else:
                    # Status tracking container with technical process feedback
                    status_container = st.status("⚗️ Distilling Architectural Signals...", expanded=True)
                    
                    try:
                        with status_container:
                            st.write("🔍 Checking local cache for existing analysis...")
                            st.write("⚗️ Initializing credit-efficient scraping protocol...")
                            st.write("🎯 Preparing targeted external signal hunt...")
                            
                            # Run the agentic triangulation loop
                            inputs, results = st.session_state.ba_assistant.run_triangulation(
                                target_url,
                                job_posting,
                                st.session_state.firecrawl_manager,
                                st.session_state.db_manager
                            )
                            
                            st.session_state.inputs = inputs
                            st.session_state.results = results
                            
                        status_container.update(
                            label="✅ Architectural Reality Check Complete!", 
                            state="complete", 
                            expanded=False
                        )
                        st.success(f"🏗️ Technical assessment complete for {inputs.company_name} - contradictions identified!")
                        st.rerun()
                            
                    except Exception as e:
                        status_container.update(
                            label="❌ Technical Analysis Failed", 
                            state="error", 
                            expanded=True
                        )
                        st.error(f"Architectural assessment failed: {str(e)}")
                        logger.error(f"Technical analysis pipeline failed: {e}")
        
        st.divider()
        
        # Prompt Management Section
        with st.expander("⚙️ Prompt Management", expanded=False):
            st.write("**Edit the system prompts used in each stage:**")
            
            prompt_tabs = st.tabs(["Extract", "Diagnose", "Hook", "Audit", "Close"])
            
            with prompt_tabs[0]:
                st.session_state.prompts["extract_signals"] = st.text_area(
                    "Signal Extraction Prompt",
                    value=st.session_state.prompts["extract_signals"],
                    height=200,
                    key="prompt_extract"
                )
            
            with prompt_tabs[1]:
                st.session_state.prompts["diagnose"] = st.text_area(
                    "Diagnostic Prompt",
                    value=st.session_state.prompts["diagnose"],
                    height=200,
                    key="prompt_diagnose"
                )
            
            with prompt_tabs[2]:
                st.session_state.prompts["generate_hook"] = st.text_area(
                    "Hook Generation Prompt",
                    value=st.session_state.prompts["generate_hook"],
                    height=200,
                    key="prompt_hook"
                )
            
            with prompt_tabs[3]:
                st.session_state.prompts["generate_audit"] = st.text_area(
                    "Audit Generation Prompt",
                    value=st.session_state.prompts["generate_audit"],
                    height=200,
                    key="prompt_audit"
                )
            
            with prompt_tabs[4]:
                st.session_state.prompts["generate_close"] = st.text_area(
                    "Close Generation Prompt",
                    value=st.session_state.prompts["generate_close"],
                    height=200,
                    key="prompt_close"
                )
            
            col_reset, col_update, col_save = st.columns(3)
            
            with col_reset:
                if st.button("🔄 Factory Reset"):
                    if 'ba_assistant' in st.session_state:
                        st.session_state.prompts = st.session_state.ba_assistant.reset_to_factory_defaults()
                        st.rerun()
            
            with col_update:
                if st.button("💾 Update Assistant"):
                    if 'ba_assistant' in st.session_state:
                        api_key = st.secrets.get("OPENAI_API_KEY")
                        st.session_state.ba_assistant = BAAssistant(api_key, st.session_state.prompts)
                        st.success("Assistant updated with new prompts!")
            
            with col_save:
                if st.button("⭐ Save as Defaults"):
                    if 'ba_assistant' in st.session_state:
                        if st.session_state.ba_assistant.save_custom_prompts(st.session_state.prompts):
                            st.success("Prompts saved as new defaults!")
                        else:
                            st.error("Failed to save prompts")
            
            st.info("💡 **Tip:** Edit prompts above, click 'Update Assistant', then 'Save as Defaults' to make your changes permanent for future sessions.")
    
    with col2:
        st.header("⚗️ Architectural Reality Check Results")
        
        if st.session_state.results and st.session_state.inputs:
            results = st.session_state.results
            inputs = st.session_state.inputs
            company_name = inputs.company_name or results.signals.get('company_name', 'Unknown Company')
            pdf_generator = PDFGenerator()
            word_generator = WordGenerator()
            timestamp = datetime.now().strftime('%Y%m%d_%H%M')
            
            # Show triangulation summary (expanded by default to show the evidence)
            with st.expander("🔍 Triangulation Evidence", expanded=True):
                if inputs.scraped_content:
                    st.subheader("📄 Company Self-Perception")
                    st.caption("Extraction Method: Firecrawl API (cloud rendered)")
                    st.markdown(inputs.scraped_content[:500] + "..." if len(inputs.scraped_content) > 500 else inputs.scraped_content)
                
                if inputs.job_posting:
                    st.subheader("💼 Job Posting Signals")
                    st.info(f"Role: {inputs.job_posting}")
                    st.caption("Job posting requirements are integrated into the technical analysis")
                
                if inputs.external_signals and "No external signals found" not in inputs.external_signals:
                    st.subheader("🔍 External Technical Signals")
                    # Count the number of signal sources
                    signal_count = len(inputs.external_signals.split("---")) if "---" in inputs.external_signals else 1
                    st.info(f"Found {signal_count} external signal sources")
                    # Show all external signals (not truncated)
                    st.markdown(inputs.external_signals)
                else:
                    st.warning("⚠️ No external signals found - analysis based on company narrative only")
            
            st.divider()
            
            # Helper function for individual downloads
            def create_download_buttons(section_title: str, content: str, file_prefix: str):
                col_md, col_pdf, col_word = st.columns(3)
                
                with col_md:
                    # Individual Markdown download
                    markdown_content = f"""# {section_title}
Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}
Company: {company_name}

{content}
"""
                    st.download_button(
                        label="📄 MD",
                        data=markdown_content,
                        file_name=f"{file_prefix}_{timestamp}.md",
                        mime="text/markdown",
                        key=f"md_{file_prefix}"
                    )
                
                with col_pdf:
                    # Individual PDF download
                    try:
                        pdf_bytes = pdf_generator.generate_section_pdf(section_title, content, company_name)
                        st.download_button(
                            label="📄 PDF",
                            data=pdf_bytes,
                            file_name=f"{file_prefix}_{timestamp}.pdf",
                            mime="application/pdf",
                            key=f"pdf_{file_prefix}"
                        )
                    except Exception as e:
                        st.error(f"PDF generation failed: {str(e)}")
                
                with col_word:
                    # Individual Word download
                    try:
                        word_bytes = word_generator.generate_section_word(section_title, content, company_name)
                        st.download_button(
                            label="📄 DOCX",
                            data=word_bytes,
                            file_name=f"{file_prefix}_{timestamp}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"word_{file_prefix}"
                        )
                    except Exception as e:
                        st.error(f"Word generation failed: {str(e)}")
            
            # Hook Section
            st.subheader("🎯 First Engagement Hook")
            st.info(results.hook)
            create_download_buttons("First Engagement Hook", results.hook, "hook")
            
            st.divider()
            
            # Diagnostic Section
            st.subheader("🔍 Diagnostic Insight")
            st.markdown(results.diagnosis)
            create_download_buttons("Diagnostic Insight", results.diagnosis, "diagnosis")
            
            st.divider()
            
            # Audit Section
            st.subheader("📋 AI Readiness Audit")
            st.markdown(results.audit)
            create_download_buttons("AI Readiness Audit", results.audit, "audit")
            
            st.divider()
            
            # Close Section
            st.subheader("🤝 Conversation Close")
            st.write(results.close)
            create_download_buttons("Conversation Close", results.close, "close")
            
            st.divider()
            
            # Complete Report Download Options
            st.subheader("💾 Complete Report")
            st.caption("Download all sections as a single file")
            
            col_md, col_pdf, col_word = st.columns(3)
            
            with col_md:
                # Complete Markdown download
                complete_markdown = f"""# AI Readiness Audit Report
Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}
Company: {company_name}

## First Engagement Hook
{results.hook}

## Diagnostic Insight  
{results.diagnosis}

## AI Readiness Audit
{results.audit}

## Conversation Close
{results.close}
"""
                st.download_button(
                    label="📄 Complete Report (MD)",
                    data=complete_markdown,
                    file_name=f"complete_ai_readiness_report_{timestamp}.md",
                    mime="text/markdown",
                    key="complete_md"
                )
            
            with col_pdf:
                # Complete PDF download
                try:
                    pdf_bytes = pdf_generator.generate_pdf(results, company_name)
                    st.download_button(
                        label="📄 Complete Report (PDF)",
                        data=pdf_bytes,
                        file_name=f"complete_ai_readiness_report_{timestamp}.pdf",
                        mime="application/pdf",
                        key="complete_pdf"
                    )
                except Exception as e:
                    st.error(f"Complete PDF generation failed: {str(e)}")
            
            with col_word:
                # Complete Word download
                try:
                    word_bytes = word_generator.generate_word(results, company_name)
                    st.download_button(
                        label="📄 Complete Report (DOCX)",
                        data=word_bytes,
                        file_name=f"complete_ai_readiness_report_{timestamp}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="complete_word"
                    )
                except Exception as e:
                    st.error(f"Complete Word generation failed: {str(e)}")
        else:
            st.info("👈 Enter a target company URL and click 'Start Architecture Analysis' to begin technical assessment.")
            
            # Show what the technical assessment will do
            st.markdown("""
            **⚗️ The Senior Architect Assessment Engine will:**
            
            1. **⚗️ Distill** company self-perception from website (credit-efficient scraping)
            2. **🔍 Hunt** for external technical contradiction signals (2-query limit)
            3. **⚖️ Verify** tech-stack contradictions vs. marketing claims
            4. **🏗️ Generate** brutal technical assessment focusing on architectural failure modes
            
            *Results cached for instant retrieval. Zero exposure outreach ready.*
            """)

if __name__ == "__main__":
    main()