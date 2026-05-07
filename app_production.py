import streamlit as st
import openai
import json
import hashlib
import logging
import time
import os
import sqlite3
from datetime import datetime
from typing import Dict, Any, Optional, List, Tuple
from dataclasses import dataclass, asdict
from urllib.parse import urlparse
import requests
from playwright.sync_api import sync_playwright
from fpdf import FPDF
import markdown
from bs4 import BeautifulSoup
import re
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
import io

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

_SCRAPE_TIMEOUT = 15000        # milliseconds (15 seconds for page load)
_SCRAPE_MAX_CHARS = 40_000    # ~10k tokens; keeps LLM context manageable

def _is_url(text: str) -> bool:
    """Return True if text looks like an http/https URL."""
    try:
        p = urlparse(text.strip())
        return p.scheme in ("http", "https") and bool(p.netloc)
    except Exception:
        return False


def discover_relevant_pages(page, base_url: str) -> List[str]:
    """
    Discover relevant pages for AI readiness analysis from a website.
    
    Returns list of URLs to scrape for comprehensive company analysis.
    """
    try:
        # Clean the input URL
        base_url = base_url.strip()
        
        # Comprehensive keywords for business analysis - covering all aspects of modern companies
        relevant_keywords = [
            # Core company info
            'about', 'team', 'company', 'people', 'leadership', 'founders', 'who-we-are', 'our-story', 'mission', 'vision',
            
            # Products & Services  
            'products', 'services', 'solutions', 'platform', 'technology', 'tech', 'software', 'tools', 'api', 'features',
            'pricing', 'plans', 'offerings', 'capabilities', 'stack', 'infrastructure',
            
            # Intelligence & Research (as you mentioned)
            'intelligence', 'research', 'insights', 'analytics', 'data', 'analysis', 'reports', 'findings',
            'expertise', 'expert', 'experts', 'opportunities', 'consulting', 'advisory', 'specialists',
            'data-engine', 'engine', 'algorithms', 'machine-learning', 'ai', 'artificial-intelligence',
            
            # Business operations
            'careers', 'jobs', 'hiring', 'work', 'culture', 'join-us', 'openings', 'positions', 'talent',
            'resources', 'support', 'help', 'documentation', 'docs', 'guides', 'tutorials',
            
            # Marketing & Content
            'blog', 'news', 'press', 'media', 'insights', 'articles', 'updates', 'announcements',
            'case-studies', 'customers', 'clients', 'portfolio', 'success', 'testimonials', 'reviews',
            'use-cases', 'examples', 'stories', 'showcase',
            
            # Business development
            'partners', 'partnerships', 'integrations', 'ecosystem', 'marketplace', 'network',
            'events', 'webinars', 'workshops', 'training', 'education', 'learning',
            
            # Contact & Legal
            'contact', 'get-in-touch', 'reach-out', 'connect', 'talk', 'demo', 'sales',
            'investor', 'investors', 'funding', 'investment', 'venture', 'equity',
            'legal', 'privacy', 'terms', 'security', 'compliance', 'governance',
            
            # Industry-specific
            'research', 'development', 'innovation', 'labs', 'r&d', 'science', 'engineering',
            'enterprise', 'business', 'corporate', 'professional', 'industrial',
            'startups', 'scale', 'growth', 'expansion', 'global', 'international',
            
            # Modern business terms
            'saas', 'cloud', 'digital', 'automation', 'workflow', 'optimization',
            'performance', 'efficiency', 'productivity', 'transformation', 'modernization'
        ]
        
        # Get all links from the page with more comprehensive detection
        links = page.evaluate("""
            () => {
                const links = [];
                document.querySelectorAll('a[href]').forEach(link => {
                    const href = link.getAttribute('href');
                    const text = (link.textContent || link.innerText || '').toLowerCase().trim();
                    const title = (link.getAttribute('title') || '').toLowerCase();
                    const ariaLabel = (link.getAttribute('aria-label') || '').toLowerCase();
                    if (href) {
                        links.push({ 
                            href, 
                            text,
                            title,
                            ariaLabel,
                            fullText: text + ' ' + title + ' ' + ariaLabel
                        });
                    }
                });
                return links;
            }
        """)
        
        logger.info(f"🔍 Found {len(links)} total links on page")
        
        # Track URLs by source
        discovered_urls = set()  # Actually found on the page
        guessed_urls = set()     # From our common paths list
        
        base_domain = urlparse(base_url).netloc.strip()
        base_scheme = urlparse(base_url).scheme.strip()
        
        logger.debug(f"🔍 Base URL: '{base_url}'")
        logger.debug(f"🔍 Base domain: '{base_domain}' (length: {len(base_domain)})")
        logger.debug(f"🔍 Base scheme: '{base_scheme}'")
        
        # Comprehensive common page patterns - includes your specific examples
        common_paths = [
            # Core company pages
            '/about', '/about-us', '/company', '/team', '/people', '/leadership', '/founders',
            '/mission', '/vision', '/who-we-are', '/our-story',
            
            # Products & Services
            '/products', '/services', '/solutions', '/platform', '/technology', '/tech', '/software',
            '/tools', '/api', '/features', '/pricing', '/plans', '/offerings', '/capabilities',
            
            # Intelligence & Research (your examples)
            '/intelligence', '/research', '/insights', '/analytics', '/data', '/analysis',
            '/expert-opportunities', '/experts', '/expertise', '/consulting', '/advisory',
            '/data-engine', '/engine', '/algorithms', '/machine-learning', '/ai',
            
            # Business operations  
            '/careers', '/jobs', '/join-us', '/hiring', '/work-with-us', '/openings',
            '/culture', '/talent', '/positions',
            
            # Content & Marketing
            '/blog', '/news', '/press', '/media', '/insights', '/articles', '/updates',
            '/case-studies', '/customers', '/clients', '/portfolio', '/success', '/testimonials',
            '/use-cases', '/examples', '/stories', '/showcase',
            
            # Business development
            '/partners', '/partnerships', '/integrations', '/ecosystem', '/marketplace',
            '/events', '/webinars', '/workshops', '/training', '/education',
            
            # Contact & Sales (your "Get in touch" example)
            '/contact', '/contact-us', '/get-in-touch', '/reach-out', '/connect',
            '/talk', '/demo', '/sales', '/support', '/help',
            
            # Investment & Legal
            '/investors', '/funding', '/investment', '/legal', '/privacy', '/terms',
            '/security', '/compliance',
            
            # R&D and Innovation
            '/research', '/development', '/innovation', '/labs', '/r-and-d', '/science',
            '/engineering', '/resources', '/documentation', '/docs', '/guides'
        ]
        
        # Add common paths as potential URLs (these are guesses)
        for path in common_paths[:3]:  # Just debug first 3
            potential_url = f"{base_scheme}://{base_domain}{path}"
            logger.debug(f"🔍 Creating URL: '{base_scheme}' + '://' + '{base_domain}' + '{path}' = '{potential_url}'")
            guessed_urls.add(potential_url)
        
        # Add the rest without debug spam
        for path in common_paths[3:]:
            potential_url = f"{base_scheme}://{base_domain}{path}"
            guessed_urls.add(potential_url)
        
        # Process discovered links
        for link in links:
            href = link['href']
            text = link['text']
            full_text = link['fullText']
            
            # Convert relative URLs to absolute
            if href.startswith('/'):
                full_url = f"{base_scheme}://{base_domain}{href}"
            elif href.startswith('http'):
                full_url = href
            elif href.startswith('#'):
                continue  # Skip anchors
            else:
                # Handle other relative URLs
                full_url = f"{base_scheme}://{base_domain}/{href}"
            
            # Clean up URL - remove extra spaces and normalize  
            full_url = full_url.replace(' /', '/').replace('/ ', '/').strip()
                
            # Check if URL belongs to same domain
            try:
                if urlparse(full_url).netloc != base_domain:
                    continue
            except:
                continue
                
            # Check if URL or link text contains relevant keywords
            url_lower = full_url.lower()
            
            if any(keyword in url_lower or keyword in full_text for keyword in relevant_keywords):
                discovered_urls.add(full_url)
                logger.info(f"📄 Found relevant link: {full_url} (matched: {text})")
                logger.debug(f"🔍 URL before cleanup: '{href}' -> after: '{full_url}'")
        
        # Debug: Show what we found
        logger.info(f"🔍 Raw discovered URLs: {list(discovered_urls)[:5]}...")
        logger.info(f"🔍 Raw guessed URLs: {list(guessed_urls)[:5]}...")
        
        # Remove the base URL from both sets
        discovered_urls.discard(base_url)
        guessed_urls.discard(base_url)
        
        # Filter out duplicate paths and common non-content pages
        filtered_discovered = []
        filtered_guessed = []
        seen_paths = set()
        
        # First process discovered URLs (they take priority)
        for url in discovered_urls:
            path = urlparse(url).path.lower()
            
            # Skip if we've seen this path or if it's a non-content page
            if path in seen_paths or any(skip in path for skip in ['/wp-', '/admin', '/login', '/register', '/search']):
                continue
                
            seen_paths.add(path)
            filtered_discovered.append(url)
        
        # Then process guessed URLs (only if we haven't seen the path)
        for url in guessed_urls:
            path = urlparse(url).path.lower()
            
            # Skip if we've seen this path or if it's a non-content page
            if path in seen_paths or any(skip in path for skip in ['/wp-', '/admin', '/login', '/register', '/search']):
                continue
                
            seen_paths.add(path)
            filtered_guessed.append(url)
        
        # Now we have clean separation
        
        logger.info(f"🎯 Discovered real pages: {len(filtered_discovered)}, Guessed paths: {len(filtered_guessed)}")
        
        # Prioritize DISCOVERED pages first (they actually exist!)
        prioritized_urls = []
        
        # Priority 1: Actually discovered pages
        priority_keywords = ['intelligence', 'research', 'data-engine', 'experts', 'about', 'team', 'products', 'careers']
        for keyword in priority_keywords:
            for url in filtered_discovered:
                if keyword in url.lower() and url not in prioritized_urls:
                    prioritized_urls.append(url)
                    logger.info(f"✅ Prioritized discovered page: {url}")
                    if len(prioritized_urls) >= 5:
                        break
            if len(prioritized_urls) >= 5:
                break
        
        # Priority 2: Fill remaining slots with other discovered pages
        for url in filtered_discovered:
            if url not in prioritized_urls:
                prioritized_urls.append(url)
                logger.info(f"➕ Added discovered page: {url}")
                if len(prioritized_urls) >= 5:
                    break
        
        # Priority 3: Only add guessed paths if we have slots left
        if len(prioritized_urls) < 5:
            logger.info(f"🔍 Trying guessed paths to fill remaining slots...")
            for url in filtered_guessed:
                if url not in prioritized_urls:
                    prioritized_urls.append(url)
                    logger.info(f"🤞 Added guessed path: {url}")
                    if len(prioritized_urls) >= 5:
                        break
        
        logger.info(f"📋 Prioritized {len(prioritized_urls)} pages for scraping: {prioritized_urls}")
        
        # Debug: Show individual URLs to check for spaces
        for i, url in enumerate(prioritized_urls):
            logger.debug(f"🔍 URL {i+1}: '{url}' (length: {len(url)})")
            
        return prioritized_urls
        
    except Exception as e:
        logger.error(f"Page discovery failed: {e}")
        return []


def test_page_existence(page, url: str) -> Tuple[bool, int, str]:
    """
    Quickly test if a page exists without full content extraction.
    Returns (exists, status_code, title).
    """
    try:
        response = page.goto(url, wait_until="domcontentloaded", timeout=10000)  # Quick test, 10s timeout
        
        if not response:
            return False, 0, ""
        
        status_code = response.status
        
        # Check if it's a successful response
        if status_code >= 200 and status_code < 400:
            # Get basic page info
            title = page.evaluate("() => document.title || 'Untitled Page'")
            
            # Quick content check - make sure it's not just an error page
            has_content = page.evaluate("""
                () => {
                    const body = document.body;
                    const text = (body?.innerText || body?.textContent || '').trim();
                    return text.length > 100; // At least 100 chars of content
                }
            """)
            
            return has_content, status_code, title
        else:
            return False, status_code, ""
            
    except Exception as e:
        logger.warning(f"⚠️ Error testing {url}: {str(e)}")
        return False, 0, ""


def scrape_single_page(page, url: str) -> Tuple[bool, str, str]:
    """
    Scrape a single page and return success, content, and page title.
    """
    try:
        response = page.goto(url, wait_until="networkidle", timeout=_SCRAPE_TIMEOUT)
        
        if not response or not response.ok:
            return False, f"Failed to load {url} - HTTP {response.status if response else 'timeout'}", ""
        
        # Get page title
        title = page.evaluate("() => document.title || 'Untitled Page'")
        
        # Execute JavaScript to clean and extract content
        content = page.evaluate("""
            () => {
                // Remove noise elements
                const noiseSelectors = [
                    'script', 'style', 'nav', 'footer', 'header', 'aside', 
                    '.navigation', '.sidebar', '.menu', '.ads', '.advertisement',
                    '.cookie-banner', '.popup', '.modal', '.overlay'
                ];
                
                noiseSelectors.forEach(selector => {
                    document.querySelectorAll(selector).forEach(el => el.remove());
                });
                
                // Find main content area
                let mainElement = 
                    document.querySelector('main') ||
                    document.querySelector('article') ||
                    document.querySelector('#content') ||
                    document.querySelector('.content') ||
                    document.querySelector('.main-content') ||
                    document.querySelector('.post-content') ||
                    document.querySelector('.entry-content') ||
                    document.querySelector('.page-content') ||
                    document.body;
                
                if (!mainElement) {
                    return 'No main content found on page';
                }
                
                // Get clean text content
                let text = mainElement.innerText || mainElement.textContent || '';
                
                // Clean up the text
                text = text
                    .replace(/\\s+/g, ' ')
                    .replace(/\\n\\s*\\n\\s*\\n+/g, '\\n\\n')
                    .trim();
                
                return text;
            }
        """)
        
        return True, content, title
        
    except Exception as e:
        return False, f"Error scraping {url}: {str(e)}", ""


def scrape_website(url: str) -> Tuple[bool, Dict[str, Any]]:
    """
    Comprehensively scrape a website including relevant pages for AI readiness analysis.
    
    Returns (success, results_dict). Results dict contains:
    - 'main_content': Combined content from all scraped pages
    - 'pages_scraped': List of dictionaries with page details
    - 'total_chars': Total characters scraped
    - 'page_count': Number of pages scraped
    """
    try:
        with sync_playwright() as p:
            # Launch browser in headless mode (no GUI)
            browser = p.chromium.launch(headless=True)
            
            # Create new page with mobile user agent to avoid bot detection
            context = browser.new_context(
                user_agent='Mozilla/5.0 (iPhone; CPU iPhone OS 14_7_1 like Mac OS X) AppleWebKit/605.1.15'
            )
            page = context.new_page()
            page.set_default_timeout(_SCRAPE_TIMEOUT)
            
            pages_scraped = []
            all_content = []
            
            # Step 1: Scrape the main page
            logger.info(f"📄 Scraping main page: {url}")
            success, content, title = scrape_single_page(page, url.strip())
            
            if not success:
                browser.close()
                return False, {"error": content}
            
            pages_scraped.append({
                "url": url,
                "title": title,
                "content": content,
                "char_count": len(content),
                "page_type": "Main Page"
            })
            all_content.append(f"=== {title} ===\nURL: {url}\n\n{content}")
            
            # Step 2: Discover ALL relevant pages
            logger.info(f"🔍 Discovering relevant pages from {url}")
            relevant_pages = discover_relevant_pages(page, url)
            logger.info(f"📋 Found {len(relevant_pages)} potential relevant pages")
            
            # Step 3: PHASE 1 - Test ALL pages for existence (quick check)
            logger.info(f"🧪 Phase 1: Testing all {len(relevant_pages)} pages for existence...")
            working_pages = []
            
            for i, page_url in enumerate(relevant_pages):
                logger.info(f"🧪 Testing {i+1}/{len(relevant_pages)}: {page_url}")
                exists, status_code, title = test_page_existence(page, page_url)
                
                if exists:
                    working_pages.append({
                        "url": page_url,
                        "title": title,
                        "status_code": status_code
                    })
                    logger.info(f"✅ Page exists: {page_url} ({status_code}) - {title}")
                else:
                    logger.info(f"❌ Page not found: {page_url} ({status_code})")
            
            logger.info(f"📊 Phase 1 complete: {len(working_pages)} working pages found out of {len(relevant_pages)} tested")
            
            # Step 4: PHASE 2 - Prioritize and scrape only WORKING pages
            if working_pages:
                logger.info(f"🎯 Phase 2: Scraping {min(5, len(working_pages))} working pages...")
                
                # Prioritize working pages
                priority_keywords = ['intelligence', 'research', 'data-engine', 'experts', 'about', 'team', 'products', 'careers']
                working_pages.sort(key=lambda p: next((i for i, keyword in enumerate(priority_keywords) 
                                                     if keyword in p['url'].lower() or keyword in p['title'].lower()), 999))
                
                # Scrape up to 5 working pages
                for i, page_info in enumerate(working_pages[:5]):
                    page_url = page_info['url']
                    try:
                        logger.info(f"📄 Scraping working page {i+1}/{min(5, len(working_pages))}: {page_url}")
                        success, content, title = scrape_single_page(page, page_url)
                        
                        if success and len(content.strip()) > 100:  # Only include substantive content
                            # Determine page type based on URL/title
                            page_type = "Other"
                            url_lower = page_url.lower()
                            title_lower = title.lower()
                            
                            if any(word in url_lower or word in title_lower for word in ['intelligence', 'ai', 'data']):
                                page_type = "Intelligence/AI"
                            elif any(word in url_lower or word in title_lower for word in ['research', 'insights', 'analysis']):
                                page_type = "Research/Insights"
                            elif any(word in url_lower or word in title_lower for word in ['expert', 'consulting', 'advisory']):
                                page_type = "Expert/Consulting"
                            elif any(word in url_lower or word in title_lower for word in ['about', 'company']):
                                page_type = "About/Company"
                            elif any(word in url_lower or word in title_lower for word in ['team', 'people', 'leadership']):
                                page_type = "Team/People"
                            elif any(word in url_lower or word in title_lower for word in ['product', 'service', 'solution']):
                                page_type = "Products/Services"
                            elif any(word in url_lower or word in title_lower for word in ['career', 'job', 'hiring']):
                                page_type = "Careers/Jobs"
                            elif any(word in url_lower or word in title_lower for word in ['blog', 'news', 'press']):
                                page_type = "Blog/News"
                            
                            pages_scraped.append({
                                "url": page_url,
                                "title": title,
                                "content": content,
                                "char_count": len(content),
                                "page_type": page_type
                            })
                            all_content.append(f"=== {title} ===\nURL: {page_url}\n\n{content}")
                            logger.info(f"📄 Successfully scraped: {page_url} ({len(content)} chars)")
                            
                    except Exception as e:
                        logger.warning(f"⚠️ Failed to scrape working page {page_url}: {str(e)}")
                        continue
            else:
                logger.warning(f"⚠️ No working pages found - all discovered pages returned errors")
            
            browser.close()
            
            # Combine all content
            combined_content = "\n\n" + "="*80 + "\n\n".join(all_content)
            
            # Apply total length limit
            if len(combined_content) > _SCRAPE_MAX_CHARS:
                combined_content = combined_content[:_SCRAPE_MAX_CHARS] + "\n\n[Content truncated — total content too long]"
            
            total_chars = len(combined_content)
            page_count = len(pages_scraped)
            
            logger.info(f"🎯 Multi-page scraping complete: {page_count} pages, {total_chars:,} total chars")
            
            return True, {
                "main_content": combined_content,
                "pages_scraped": pages_scraped,
                "total_chars": total_chars,
                "page_count": page_count
            }
            
    except Exception as exc:
        error_msg = str(exc)
        
        # Provide specific error messages for common issues
        if "timeout" in error_msg.lower():
            return False, {"error": f"Page took too long to load: {url} (timeout after 15s)"}
        elif "net::" in error_msg:
            return False, {"error": f"Network error accessing {url} - check URL and internet connection"}
        elif "browser" in error_msg.lower():
            return False, {"error": f"Browser error - Playwright may not be properly installed: {error_msg}"}
        else:
            return False, {"error": f"Scraping error for {url}: {error_msg}"}

@dataclass
class CompanyInputs:
    linkedin_url: str
    website: str        # original URL or free-text summary
    job_posting: str
    website_content: Optional[str] = None  # scraped Markdown; None when website is free-text or scrape failed

    @property
    def website_context(self) -> str:
        """Return scraped Markdown if available, otherwise the raw website field."""
        return self.website_content if self.website_content else self.website

@dataclass
class PipelineResults:
    signals: Dict[str, Any]
    diagnosis: str
    hook: str
    audit: str
    close: str
    patterns: Optional[Dict[str, Any]] = None

# ==================== MEMORY LAYER ====================

class CompanyMemory:
    """SQLite-based memory layer for company data storage"""
    
    def __init__(self, db_path: str = "companies.db"):
        self.db_path = db_path
        self.init_database()
    
    def init_database(self):
        """Initialize SQLite database with required tables"""
        with sqlite3.connect(self.db_path) as conn:
            # Companies table
            conn.execute('''
                CREATE TABLE IF NOT EXISTS companies (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    company_name TEXT NOT NULL,
                    raw_inputs TEXT NOT NULL,
                    extracted_signals TEXT NOT NULL,
                    diagnostic_summary TEXT,
                    system_classification TEXT,
                    failure_modes TEXT,
                    ai_maturity TEXT,
                    timestamp TEXT NOT NULL,
                    input_hash TEXT UNIQUE NOT NULL,
                    extended_data TEXT DEFAULT "{}"
                )
            ''')
            
            # Interaction states table
            conn.execute('''
                CREATE TABLE IF NOT EXISTS interaction_states (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    company_name TEXT NOT NULL,
                    state TEXT NOT NULL,
                    updated_at TEXT NOT NULL,
                    UNIQUE(company_name)
                )
            ''')
            
            # Cache table
            conn.execute('''
                CREATE TABLE IF NOT EXISTS cache (
                    input_hash TEXT PRIMARY KEY,
                    results TEXT NOT NULL,
                    created_at TEXT NOT NULL
                )
            ''')
            
            # Add extended_data column for backward compatibility
            try:
                conn.execute('ALTER TABLE companies ADD COLUMN extended_data TEXT DEFAULT "{}"')
            except sqlite3.OperationalError:
                # Column already exists or other schema issue
                pass
            
            conn.commit()
    
    def save_company(self, company_name: str, raw_inputs: CompanyInputs, 
                    signals: Dict[str, Any], diagnosis: str, 
                    classification: str, failure_modes: List[str], 
                    ai_maturity: str, full_results = None) -> bool:
        """Save company analysis to database with full pipeline results"""
        try:
            input_hash = self._generate_input_hash(raw_inputs)
            timestamp = datetime.now().isoformat()
            
            # Prepare extended data storage
            extended_data = {
                'hook': full_results.hook if full_results else '',
                'audit': full_results.audit if full_results else '',
                'close': full_results.close if full_results else '',
                'patterns': full_results.patterns if full_results else {}
            }
            
            with sqlite3.connect(self.db_path) as conn:
                conn.execute('''
                    INSERT OR REPLACE INTO companies 
                    (company_name, raw_inputs, extracted_signals, diagnostic_summary,
                     system_classification, failure_modes, ai_maturity, timestamp, input_hash, extended_data)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    company_name,
                    json.dumps(asdict(raw_inputs)),
                    json.dumps(signals),
                    diagnosis,
                    classification,
                    json.dumps(failure_modes),
                    ai_maturity,
                    timestamp,
                    input_hash,
                    json.dumps(extended_data)
                ))
                conn.commit()
            return True
        except Exception as e:
            logger.error(f"Failed to save company: {e}")
            return False
    
    def get_company(self, company_name: str) -> Optional[Dict[str, Any]]:
        """Retrieve company data by name with backward compatibility"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.execute(
                    'SELECT * FROM companies WHERE company_name = ? ORDER BY timestamp DESC LIMIT 1',
                    (company_name,)
                )
                row = cursor.fetchone()
                
                if row:
                    columns = [desc[0] for desc in cursor.description]
                    company_data = dict(zip(columns, row))
                    
                    # Parse JSON fields
                    company_data['raw_inputs'] = json.loads(company_data['raw_inputs'])
                    company_data['extracted_signals'] = json.loads(company_data['extracted_signals'])
                    company_data['failure_modes'] = json.loads(company_data['failure_modes'])
                    
                    # Handle extended data with backward compatibility
                    if 'extended_data' in company_data and company_data['extended_data']:
                        try:
                            extended = json.loads(company_data['extended_data'])
                            company_data.update(extended)
                        except (json.JSONDecodeError, TypeError):
                            # Legacy record without extended data
                            company_data['hook'] = 'Legacy record - hook not stored'
                            company_data['audit'] = 'Legacy record - audit not stored'
                            company_data['close'] = 'Legacy record - close not stored'
                            company_data['patterns'] = {}
                    else:
                        # Legacy record
                        company_data['hook'] = 'Legacy record - hook not stored'
                        company_data['audit'] = 'Legacy record - audit not stored' 
                        company_data['close'] = 'Legacy record - close not stored'
                        company_data['patterns'] = {}
                    
                    return company_data
        except Exception as e:
            logger.error(f"Failed to get company {company_name}: {e}")
        return None
    
    def list_companies(self, limit: int = 50) -> List[Dict[str, Any]]:
        """List all companies with basic info"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.execute('''
                    SELECT company_name, system_classification, ai_maturity, timestamp
                    FROM companies 
                    ORDER BY timestamp DESC 
                    LIMIT ?
                ''', (limit,))
                
                columns = [desc[0] for desc in cursor.description]
                return [dict(zip(columns, row)) for row in cursor.fetchall()]
        except Exception as e:
            logger.error(f"Failed to list companies: {e}")
            return []
    
    def _generate_input_hash(self, inputs: CompanyInputs) -> str:
        """Generate hash from input combination"""
        combined = f"{inputs.linkedin_url}{inputs.website}{inputs.job_posting}"
        return hashlib.md5(combined.encode()).hexdigest()

# ==================== PATTERN DETECTION ENGINE ====================

class PatternEngine:
    """Cross-company pattern detection and similarity analysis"""
    
    def __init__(self, memory: CompanyMemory):
        self.memory = memory
    
    def detect_patterns(self, signals: Dict[str, Any], company_name: str) -> Dict[str, Any]:
        """Detect patterns across similar companies"""
        similar_companies = self._find_similar_companies(signals, company_name)
        recurring_failures = self._get_recurring_failure_modes(similar_companies)
        common_patterns = self._extract_common_patterns(similar_companies, signals)
        
        return {
            "similar_companies": similar_companies,
            "recurring_failure_modes": recurring_failures,
            "common_patterns": common_patterns
        }
    
    def _find_similar_companies(self, target_signals: Dict[str, Any], exclude_name: str) -> List[Dict[str, Any]]:
        """Find top 3 most similar companies based on tech stack and architecture"""
        companies = self.memory.list_companies()
        similarities = []
        
        target_tech_stack = set(target_signals.get('tech_stack', []))
        target_architecture = set(target_signals.get('architecture_keywords', []))
        
        for company in companies:
            if company['company_name'] == exclude_name:
                continue
                
            company_data = self.memory.get_company(company['company_name'])
            if not company_data:
                continue
                
            company_signals = company_data['extracted_signals']
            company_tech = set(company_signals.get('tech_stack', []))
            company_arch = set(company_signals.get('architecture_keywords', []))
            
            # Calculate similarity scores
            tech_overlap = len(target_tech_stack.intersection(company_tech))
            arch_overlap = len(target_architecture.intersection(company_arch))
            
            # Weight tech stack more heavily than architecture keywords
            similarity_score = (tech_overlap * 2) + arch_overlap
            
            if similarity_score > 0:
                similarities.append({
                    'company_name': company['company_name'],
                    'similarity_score': similarity_score,
                    'tech_overlap': list(target_tech_stack.intersection(company_tech)),
                    'arch_overlap': list(target_architecture.intersection(company_arch)),
                    'classification': company['system_classification'],
                    'ai_maturity': company['ai_maturity']
                })
        
        # Return top 3 most similar
        similarities.sort(key=lambda x: x['similarity_score'], reverse=True)
        return similarities[:3]
    
    def _get_recurring_failure_modes(self, similar_companies: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Get ranked recurring failure modes from similar companies"""
        failure_count = {}
        
        for similar in similar_companies:
            company_data = self.memory.get_company(similar['company_name'])
            if company_data and company_data['failure_modes']:
                for failure in company_data['failure_modes']:
                    failure_count[failure] = failure_count.get(failure, 0) + 1
        
        # Convert to ranked list
        ranked_failures = [
            {"failure_mode": failure, "frequency": count}
            for failure, count in sorted(failure_count.items(), key=lambda x: x[1], reverse=True)
        ]
        
        return ranked_failures[:5]  # Top 5
    
    def _extract_common_patterns(self, similar_companies: List[Dict[str, Any]], target_signals: Dict[str, Any]) -> str:
        """Extract common architectural patterns and blind spots"""
        if not similar_companies:
            return "No similar companies in database for pattern analysis."
        
        # Analyze classification patterns
        classifications = [comp['classification'] for comp in similar_companies if comp['classification']]
        ai_maturities = [comp['ai_maturity'] for comp in similar_companies if comp['ai_maturity']]
        
        # Common tech patterns
        all_tech = []
        for similar in similar_companies:
            company_data = self.memory.get_company(similar['company_name'])
            if company_data:
                all_tech.extend(company_data['extracted_signals'].get('tech_stack', []))
        
        tech_frequency = {}
        for tech in all_tech:
            tech_frequency[tech] = tech_frequency.get(tech, 0) + 1
        
        common_tech = [tech for tech, freq in tech_frequency.items() if freq >= 2]
        
        # Generate synthesis
        if classifications:
            most_common_class = max(set(classifications), key=classifications.count) if classifications else "Unknown"
            
            pattern_summary = f"Companies with similar tech stacks typically classified as {most_common_class}. "
            if common_tech:
                pattern_summary += f"Common technologies: {', '.join(common_tech[:3])}. "
            if len(similar_companies) >= 2:
                pattern_summary += f"Pattern observed across {len(similar_companies)} similar companies in database."
            else:
                pattern_summary += "Limited pattern data - early analysis."
        else:
            pattern_summary = "Insufficient historical data for reliable pattern analysis."
        
        return pattern_summary

# ==================== INTERACTION STATE LAYER ====================

class InteractionState:
    """Track interaction states per company"""
    
    STATES = [
        'HOOK_SENT',
        'ENGAGED', 
        'CLARIFYING',
        'QUALIFIED',
        'AUDIT_SHARED',
        'DROPPED'
    ]
    
    def __init__(self, memory: CompanyMemory):
        self.memory = memory
    
    def update_state(self, company_name: str, state: str) -> bool:
        """Update interaction state for a company"""
        if state not in self.STATES:
            logger.error(f"Invalid state: {state}")
            return False
            
        try:
            with sqlite3.connect(self.memory.db_path) as conn:
                conn.execute('''
                    INSERT OR REPLACE INTO interaction_states 
                    (company_name, state, updated_at)
                    VALUES (?, ?, ?)
                ''', (company_name, state, datetime.now().isoformat()))
                conn.commit()
            return True
        except Exception as e:
            logger.error(f"Failed to update state: {e}")
            return False
    
    def get_state(self, company_name: str) -> Optional[str]:
        """Get current interaction state for a company"""
        try:
            with sqlite3.connect(self.memory.db_path) as conn:
                cursor = conn.execute(
                    'SELECT state FROM interaction_states WHERE company_name = ?',
                    (company_name,)
                )
                row = cursor.fetchone()
                return row[0] if row else None
        except Exception as e:
            logger.error(f"Failed to get state: {e}")
            return None

# ==================== LIGHTWEIGHT CACHING ====================

class LightweightCache:
    """Hash-based caching for pipeline results"""
    
    def __init__(self, memory: CompanyMemory):
        self.memory = memory
    
    def get(self, inputs: CompanyInputs) -> Optional[PipelineResults]:
        """Get cached results if available"""
        try:
            input_hash = self._generate_hash(inputs)
            
            with sqlite3.connect(self.memory.db_path) as conn:
                cursor = conn.execute(
                    'SELECT results FROM cache WHERE input_hash = ?',
                    (input_hash,)
                )
                row = cursor.fetchone()
                
                if row:
                    cached_data = json.loads(row[0])
                    return PipelineResults(**cached_data)
        except Exception as e:
            logger.error(f"Cache retrieval failed: {e}")
        return None
    
    def set(self, inputs: CompanyInputs, results: PipelineResults):
        """Cache pipeline results"""
        try:
            input_hash = self._generate_hash(inputs)
            
            with sqlite3.connect(self.memory.db_path) as conn:
                conn.execute('''
                    INSERT OR REPLACE INTO cache 
                    (input_hash, results, created_at)
                    VALUES (?, ?, ?)
                ''', (
                    input_hash,
                    json.dumps(asdict(results)),
                    datetime.now().isoformat()
                ))
                conn.commit()
        except Exception as e:
            logger.error(f"Cache storage failed: {e}")
    
    def _generate_hash(self, inputs: CompanyInputs) -> str:
        """Generate hash from inputs"""
        combined = f"{inputs.linkedin_url}{inputs.website}{inputs.job_posting}"
        return hashlib.md5(combined.encode()).hexdigest()

# ==================== ADVISORY FLOW HELPERS ====================

class AdvisoryFlows:
    """Helper methods for advisory conversation flows"""
    
    def __init__(self, ba_assistant):
        self.ba_assistant = ba_assistant
    
    def generate_hook(self, signals: Dict[str, Any]) -> str:
        """Generate initial hook based on signals"""
        return self.ba_assistant._make_llm_call(
            self.ba_assistant.prompts["generate_hook"],
            f"Generate hook for signals: {json.dumps(signals, indent=2)}"
        )
    
    def generate_followup_insight(self, signals: Dict[str, Any]) -> str:
        """Generate follow-up insight for engaged prospects"""
        followup_prompt = """You are providing a follow-up technical insight to a founder who responded to your initial hook.

Requirements:
- 2-3 sentences maximum
- Reference a specific technical constraint they likely haven't considered
- No sales language, just peer-to-peer technical observation
- Build on the initial contradiction you identified

Tone: Technical peer sharing an additional observation"""
        
        return self.ba_assistant._make_llm_call(
            followup_prompt,
            f"Generate followup insight for engaged founder with signals: {json.dumps(signals, indent=2)}"
        )
    
    def generate_clarification_response(self, signals: Dict[str, Any]) -> str:
        """Generate response to clarifying questions"""
        clarification_prompt = """You are responding to clarifying questions from a founder about your initial technical observation.

Requirements:
- Be specific about the technical constraint you identified
- Provide 1-2 concrete examples of how this typically manifests
- Keep it conversational and helpful
- No consulting pitch, just technical clarity

Tone: Technical peer explaining your reasoning"""
        
        return self.ba_assistant._make_llm_call(
            clarification_prompt,
            f"Clarify technical observation for signals: {json.dumps(signals, indent=2)}"
        )
    
    def generate_audit(self, signals: Dict[str, Any], patterns: Dict[str, Any]) -> str:
        """Generate enhanced audit with pattern insights"""
        enhanced_audit_prompt = self.ba_assistant.prompts["generate_audit"] + f"""

ADDITIONAL CONTEXT - Cross-Company Patterns:
{json.dumps(patterns, indent=2)}

Include a new section at the end:

## Systemic Patterns Across Similar Companies
Based on analysis of similar companies in our database:

- Top 3 recurring failure modes from companies with similar tech stacks
- Common architectural blind spots observed across this pattern
- Brief synthesis (3-4 lines, specific not generic)

Reference the pattern data provided above."""

        return self.ba_assistant._make_llm_call(
            enhanced_audit_prompt,
            f"Generate enhanced audit with patterns for signals: {json.dumps(signals, indent=2)}"
        )

# ==================== PDF GENERATOR ====================

class PDFGenerator:
    def __init__(self):
        pass
    
    def _clean_text(self, text: str) -> str:
        """Remove markdown and clean text for PDF"""
        html = markdown.markdown(text)
        soup = BeautifulSoup(html, 'html.parser')
        cleaned = soup.get_text()
        cleaned = re.sub(r'\n\s*\n', '\n\n', cleaned)
        cleaned = cleaned.replace('"', '"').replace('"', '"')
        cleaned = cleaned.replace(''', "'").replace(''', "'")
        cleaned = cleaned.replace('–', '-').replace('—', '-')
        cleaned = cleaned.replace('…', '...')
        cleaned = ''.join(char if ord(char) < 128 else '?' for char in cleaned)
        return cleaned.strip()
    
    def _add_content_to_pdf(self, pdf: FPDF, title: str, content: str, company_name: str = None):
        """Add content section to PDF"""
        pdf.set_font('Times', 'B', 16)
        pdf.cell(0, 10, f'IDPETECH - BA Assistant: {title}', ln=True, align='C')
        pdf.set_font('Times', '', 10)
        pdf.cell(0, 10, f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}', ln=True, align='C')
        pdf.ln(5)
        
        if company_name:
            pdf.set_font('Times', 'B', 14)
            pdf.cell(0, 10, f'Company: {company_name}', ln=True)
            pdf.ln(5)
        
        pdf.set_font('Times', '', 10)
        cleaned_content = self._clean_text(content)
        
        lines = cleaned_content.split('\n')
        for line in lines:
            if len(line) > 80:
                words = line.split(' ')
                current_line = ""
                for word in words:
                    if len(current_line + word) < 80:
                        current_line += word + " "
                    else:
                        if current_line:
                            pdf.cell(0, 5, current_line.strip(), ln=True)
                        current_line = word + " "
                if current_line:
                    pdf.cell(0, 5, current_line.strip(), ln=True)
            else:
                pdf.cell(0, 5, line, ln=True)
    
    def generate_section_pdf(self, title: str, content: str, company_name: str = None) -> bytes:
        """Generate PDF for individual section"""
        pdf = FPDF()
        pdf.add_page()
        self._add_content_to_pdf(pdf, title, content, company_name)
        return bytes(pdf.output())
    
    def generate_pdf(self, results: PipelineResults, company_name: str) -> bytes:
        """Generate complete PDF report"""
        pdf = FPDF()
        pdf.add_page()
        
        pdf.set_font('Times', 'B', 16)
        pdf.cell(0, 10, 'IDPETECH - BA Assistant: Complete AI Readiness Report', ln=True, align='C')
        pdf.set_font('Times', '', 10)
        pdf.cell(0, 10, f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}', ln=True, align='C')
        pdf.ln(5)
        
        pdf.set_font('Times', 'B', 14)
        pdf.cell(0, 10, f'Company: {company_name}', ln=True)
        pdf.ln(5)
        
        sections = [
            ("First Engagement Hook", results.hook),
            ("Diagnostic Insight", results.diagnosis),
            ("AI Readiness Audit", results.audit),
            ("Conversation Close", results.close)
        ]
        
        for title, content in sections:
            pdf.set_font('Times', 'B', 12)
            pdf.cell(0, 10, title, ln=True)
            pdf.ln(2)
            
            pdf.set_font('Times', '', 10)
            cleaned_content = self._clean_text(content)
            
            lines = cleaned_content.split('\n')
            for line in lines:
                if len(line) > 80:
                    words = line.split(' ')
                    current_line = ""
                    for word in words:
                        if len(current_line + word) < 80:
                            current_line += word + " "
                        else:
                            if current_line:
                                pdf.cell(0, 5, current_line.strip(), ln=True)
                            current_line = word + " "
                    if current_line:
                        pdf.cell(0, 5, current_line.strip(), ln=True)
                else:
                    pdf.cell(0, 5, line, ln=True)
            pdf.ln(5)
        
        return bytes(pdf.output())

# ==================== WORD GENERATOR ====================

class WordGenerator:
    def __init__(self):
        pass
    
    def _clean_text_for_word(self, text: str) -> str:
        """Clean and prepare text for Word document"""
        html = markdown.markdown(text)
        soup = BeautifulSoup(html, 'html.parser')
        cleaned = soup.get_text()
        cleaned = re.sub(r'\n\s*\n', '\n\n', cleaned)
        return cleaned.strip()
    
    def _add_markdown_content(self, doc: Document, content: str):
        """Add markdown content to Word document with proper formatting"""
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
                
            if line.startswith('# '):
                heading = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                heading = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                heading = doc.add_heading(line[4:], level=3)
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
            elif line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)
            
            i += 1
    
    def generate_section_word(self, title: str, content: str, company_name: str = None) -> bytes:
        """Generate Word document for individual section"""
        doc = Document()
        
        header_p = doc.add_heading(f'IDPETECH - BA Assistant: {title}', level=1)
        
        meta_p = doc.add_paragraph()
        meta_p.add_run(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}\n').italic = True
        if company_name:
            meta_p.add_run(f'Company: {company_name}').italic = True
        
        doc.add_paragraph()
        
        self._add_markdown_content(doc, content)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def generate_word(self, results: PipelineResults, company_name: str) -> bytes:
        """Generate complete Word document report"""
        doc = Document()
        
        title = doc.add_heading('IDPETECH - BA Assistant: Complete AI Readiness Report', level=1)
        
        meta_p = doc.add_paragraph()
        meta_p.add_run(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}\n').italic = True
        meta_p.add_run(f'Company: {company_name}').italic = True
        
        doc.add_page_break()
        
        sections = [
            ("First Engagement Hook", results.hook),
            ("Diagnostic Insight", results.diagnosis),
            ("AI Readiness Audit", results.audit),
            ("Conversation Close", results.close)
        ]
        
        for title, content in sections:
            doc.add_heading(title, level=2)
            doc.add_paragraph()
            self._add_markdown_content(doc, content)
            doc.add_page_break()
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

# ==================== MAIN BA ASSISTANT (ORCHESTRATOR) ====================

class BAAssistant:
    def __init__(self, api_key: str, prompts: Dict[str, str] = None):
        self.client = openai.OpenAI(api_key=api_key)
        self.prompts = prompts or self._load_prompts()
        
        # Initialize production components
        self.memory = CompanyMemory()
        self.pattern_engine = PatternEngine(self.memory)
        self.interaction_state = InteractionState(self.memory)
        self.cache = LightweightCache(self.memory)
        self.advisory_flows = AdvisoryFlows(self)
    
    def _get_default_prompts(self) -> Dict[str, str]:
        """Default system prompts for each stage"""
        return {
            "extract_signals": """You are a signal extraction engine. Extract ONLY structured data from company inputs.

COMPANY NAME EXTRACTION RULES:
- From LinkedIn URLs: Extract company name from "linkedin.com/company/[name]" or "linkedin.com/in/[name]"
- From Website/Company text: Extract the primary company name mentioned
- From Job Postings: Look for "at [Company]", "[Company] is hiring", company signatures
- If multiple names found, prioritize the most formal/complete version
- Clean up: Remove "Inc", "LLC", "Ltd" unless part of brand (e.g. "OpenAI, Inc." → "OpenAI")

Return valid JSON with these fields:
{
  "company_name": "",
  "industry": "",
  "stage": "",
  "tech_stack": [],
  "role_being_hired": "",
  "role_seniority": "",
  "specific_skills_required": [],
  "business_model_signals": [],
  "scale_indicators": [],
  "ai_mentions": [],
  "technical_complexity_signals": [],
  "architecture_keywords": [],
  "data_flow_indicators": [],
  "scaling_evidence": []
}

NO interpretation. NO analysis. ONLY data extraction. Focus on accurate company_name extraction.""",
            
            "diagnose": """You are a systems architect performing a reality check on AI ambitions.

REQUIRED OUTPUT STRUCTURE:

**What They Think They're Building:**
[Extract their self-perception from job posts, marketing language, and stated goals]

**What They're Actually Building:**
[Based on technical signals - stack, architecture, hiring patterns]

**The Gap:**
[Specific contradictions between perception and technical reality]

**Hidden Bottlenecks:**
- Data architecture constraints they haven't considered
- Technical debt that will block AI implementation
- Operational complexity they're underestimating
- Team capability gaps relative to AI ambitions

**Scaling Failure Points:**
[Specific technical chokepoints that will break at 10x scale]

**AI System Classification:**
- AI-Washed: Marketing AI but building traditional systems
- AI-Assisted: Adding AI features to existing products  
- AI-Native: Core product requires AI to function
- Non-AI: Traditional software with no AI dependency

**Classification Reasoning:**
[Justify classification with specific evidence]

NO BUZZWORDS. Be surgically precise about technical realities.""",
            
            "generate_hook": """You are writing a direct, founder-to-founder message that exposes a tension.

REQUIREMENTS:
- 3-5 sentences maximum
- Identify one contradiction between what they claim and what they're actually building
- Frame as genuine curiosity, not accusation
- Reference specific technical signals from their hiring/stack
- End with a simple, non-sales question

TONE: Technical peer, not consultant. Direct but respectful.

AVOID:
- Any buzzwords (leverage, unlock, transform, scale, optimize)
- Generic AI statements
- Sales language
- Multiple topics

FORMULA:
1. Observe their stated direction
2. Note technical contradiction 
3. Express curiosity about the gap
4. Simple question

Example style: "Noticed you're hiring for X and talking about Y, but your stack suggests Z. Curious how you're thinking about that gap - is there something I'm missing about the architecture?"
""",
            
            "generate_audit": """Generate a mechanism-driven AI Readiness Audit. Focus on what will break, not what sounds good.

# AI Readiness Audit

## Executive Summary
- Current AI Classification: [AI-Washed/AI-Assisted/AI-Native/Non-AI]
- Reality Check Score: X/10 (gap between claims and capability)
- Primary Constraint: [The one thing that will break first]

## System Reality
What they actually built vs what they claim to be building.

## Constraint Analysis
**Data Architecture:** [Specific data flow bottlenecks]
**Compute Infrastructure:** [Actual processing limitations]  
**Team Capabilities:** [Skills gaps that will cause failures]
**Technical Debt:** [Legacy constraints blocking AI implementation]
**Operational Maturity:** [Process gaps at scale]

## Failure Mode Predictions
The top 3 ways their AI ambitions will fail, with specific technical triggers.

## Contrarian Insights
**What everyone else is telling them:** [Common advice they're getting]
**What's actually true:** [Uncomfortable reality they need to face]
**The bet they're really making:** [What has to be true for them to succeed]

## Mechanism-Based Recommendations
Specific technical interventions, not strategic advice:
- [Concrete system changes with clear outcomes]
- [Process modifications with measurable effects]
- [Architecture decisions with failure prevention]

NO BUZZWORDS. NO "BEST PRACTICES." Focus on what breaks systems and how to prevent it.""",
            
            "generate_close": """Write 2-3 sentences that surface one uncomfortable truth about their approach.

Requirements:
- Reference a specific technical tension from the analysis
- No questions, no CTAs, no next steps
- End with a direct statement about what they're actually building vs what they think they're building

Avoid: consulting language, encouragement, positivity, solutions

Style: Technical peer delivering honest feedback"""
        }
    
    def _load_prompts(self) -> Dict[str, str]:
        """Load prompts from custom file if exists, otherwise use defaults"""
        custom_prompts_file = "custom_prompts.json"
        if os.path.exists(custom_prompts_file):
            try:
                with open(custom_prompts_file, 'r') as f:
                    return json.load(f)
            except Exception as e:
                logger.warning(f"Could not load custom prompts: {e}")
        return self._get_default_prompts()
    
    def save_custom_prompts(self, prompts: Dict[str, str]) -> bool:
        """Save custom prompts as new defaults"""
        try:
            with open("custom_prompts.json", 'w') as f:
                json.dump(prompts, f, indent=2)
            return True
        except Exception as e:
            logger.error(f"Could not save custom prompts: {e}")
            return False
    
    def reset_to_factory_defaults(self) -> Dict[str, str]:
        """Reset to factory defaults and remove custom file"""
        try:
            if os.path.exists("custom_prompts.json"):
                os.remove("custom_prompts.json")
        except Exception as e:
            logger.warning(f"Could not remove custom prompts file: {e}")
        return self._get_default_prompts()
    
    def _make_llm_call(self, system_prompt: str, user_prompt: str, model: str = "gpt-4o", max_retries: int = 3) -> str:
        """Reusable LLM wrapper with retry handling"""
        for attempt in range(max_retries):
            try:
                response = self.client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    temperature=0.7
                )
                return response.choices[0].message.content
            except Exception as e:
                logger.error(f"LLM call attempt {attempt + 1} failed: {e}")
                if attempt == max_retries - 1:
                    raise e
                time.sleep(2 ** attempt)
    
    def _extract_company_from_linkedin(self, linkedin_input: str) -> str:
        """Helper to extract company name from LinkedIn URL patterns"""
        import re
        
        if not linkedin_input:
            return ""
        
        # Known company name mappings for proper capitalization
        known_companies = {
            'openai': 'OpenAI',
            'anthropic': 'Anthropic', 
            'anthropic-ai': 'Anthropic',
            'meta': 'Meta',
            'meta-ai': 'Meta AI',
            'google': 'Google',
            'microsoft': 'Microsoft',
            'amazon': 'Amazon',
            'apple': 'Apple',
            'netflix': 'Netflix',
            'uber': 'Uber',
            'airbnb': 'Airbnb',
            'salesforce': 'Salesforce',
            'shopify': 'Shopify',
            'stripe': 'Stripe',
            'figma': 'Figma',
            'notion': 'Notion',
            'slack': 'Slack',
            'zoom': 'Zoom',
            'tesla': 'Tesla',
            'spacex': 'SpaceX',
            'nvidia': 'NVIDIA',
            'amd': 'AMD',
            'intel': 'Intel',
            'ibm': 'IBM'
        }
            
        # Pattern 1: linkedin.com/company/company-name
        company_match = re.search(r'linkedin\.com/company/([^/?]+)', linkedin_input, re.IGNORECASE)
        if company_match:
            company_slug = company_match.group(1).lower()
            
            # Check for exact known company match first
            if company_slug in known_companies:
                return known_companies[company_slug]
            
            # Convert slug to readable name with smart capitalization
            name = company_slug.replace('-', ' ').replace('_', ' ')
            
            # Handle special cases for AI, API, etc.
            words = []
            for word in name.split():
                if word.lower() in ['ai', 'api', 'ui', 'ux', 'ceo', 'cto', 'cfo']:
                    words.append(word.upper())
                elif word.lower() in ['inc', 'llc', 'ltd', 'corp']:
                    words.append(word.capitalize())
                else:
                    words.append(word.capitalize())
            
            return ' '.join(words)
        
        # Pattern 2: linkedin.com/in/person-name (less reliable but try)
        person_match = re.search(r'linkedin\.com/in/([^/?]+)', linkedin_input, re.IGNORECASE)
        if person_match:
            # This is a person profile, not a company - return empty to let LLM handle it
            return ""
            
        # Pattern 3: Just company name provided as text
        if not linkedin_input.startswith('http') and len(linkedin_input.strip()) < 100:
            return linkedin_input.strip()
            
        return ""
    
    def extract_signals(self, inputs: CompanyInputs) -> Dict[str, Any]:
        """Stage 1: Extract structured signals without interpretation"""
        system_prompt = self.prompts["extract_signals"]
        
        # Pre-extract company name from LinkedIn if possible
        suggested_company = self._extract_company_from_linkedin(inputs.linkedin_url)
        
        user_prompt = f"""Extract signals from:

LinkedIn/Company: {inputs.linkedin_url}
Website Content:
{inputs.website_context}
Job Posting: {inputs.job_posting}

{f"HINT: Detected company name from LinkedIn URL: '{suggested_company}'" if suggested_company else ""}

Focus especially on extracting the correct company_name field."""

        response = self._make_llm_call(system_prompt, user_prompt)
        logger.info(f"Signals extracted: {response[:200]}...")
        
        try:
            signals = json.loads(response)
            
            # Fallback: If LLM didn't extract company name but we detected one from LinkedIn
            if not signals.get('company_name') and suggested_company:
                signals['company_name'] = suggested_company
                logger.info(f"Applied fallback company name: {suggested_company}")
            
            return signals
        except json.JSONDecodeError:
            # Even on JSON error, try to return basic structure with detected company
            fallback_signals = {
                "company_name": suggested_company or "Unknown Company",
                "error": "Failed to parse signals", 
                "raw_response": response
            }
            return fallback_signals
    
    def diagnose(self, signals: Dict[str, Any], inputs: CompanyInputs) -> str:
        """Stage 2: Strategic diagnostic reasoning"""
        system_prompt = self.prompts["diagnose"]

        user_prompt = f"""Diagnose this company:

Extracted Signals: {json.dumps(signals, indent=2)}

Original Context:
LinkedIn/Company: {inputs.linkedin_url}
Website Content:
{inputs.website_context}
Job Posting: {inputs.job_posting}"""

        response = self._make_llm_call(system_prompt, user_prompt)
        logger.info(f"Diagnosis completed: {response[:200]}...")
        return response
    
    def generate_hook(self, signals: Dict[str, Any], diagnosis: str) -> str:
        """Stage 3: Generate founder-facing outbound hook"""
        return self.advisory_flows.generate_hook(signals)
    
    def generate_audit(self, signals: Dict[str, Any], diagnosis: str, patterns: Dict[str, Any] = None) -> str:
        """Stage 4: Generate structured audit report with pattern insights"""
        if patterns:
            return self.advisory_flows.generate_audit(signals, patterns)
        else:
            system_prompt = self.prompts["generate_audit"]
            user_prompt = f"""Generate audit for:

Signals: {json.dumps(signals, indent=2)}
Diagnosis: {diagnosis}"""
            response = self._make_llm_call(system_prompt, user_prompt)
            logger.info("Audit generated")
            return response
    
    def generate_close(self, signals: Dict[str, Any], audit: str) -> str:
        """Stage 5: Generate conversation close with soft CTA"""
        system_prompt = self.prompts["generate_close"]

        user_prompt = f"""Create close based on:

Company Signals: {json.dumps(signals, indent=2)}
Audit Summary: {audit[:500]}..."""

        response = self._make_llm_call(system_prompt, user_prompt)
        logger.info("Close generated")
        return response
    
    def run_full_pipeline(self, inputs: CompanyInputs) -> PipelineResults:
        """Execute complete production pipeline with memory and pattern detection"""
        # Check cache first
        cached_result = self.cache.get(inputs)
        if cached_result:
            logger.info("Returning cached results")
            return cached_result
        
        # Stage 1: Extract signals
        signals = self.extract_signals(inputs)
        company_name = signals.get('company_name', 'Unknown Company')
        
        # Stage 2: Diagnose
        diagnosis = self.diagnose(signals, inputs)
        
        # Extract failure modes and classification from diagnosis for storage
        failure_modes = self._extract_failure_modes(diagnosis)
        classification = self._extract_classification(diagnosis)
        ai_maturity = signals.get('ai_mentions', [])
        
        # Stage 3: Pattern detection
        patterns = self.pattern_engine.detect_patterns(signals, company_name)
        
        # Stage 4: Generate hook
        hook = self.generate_hook(signals, diagnosis)
        
        # Stage 5: Generate enhanced audit with patterns
        audit = self.generate_audit(signals, diagnosis, patterns)
        
        # Stage 6: Generate close
        close = self.generate_close(signals, audit)
        
        results = PipelineResults(
            signals=signals,
            diagnosis=diagnosis,
            hook=hook,
            audit=audit,
            close=close,
            patterns=patterns
        )
        
        # Save to memory with full results
        self.memory.save_company(
            company_name, inputs, signals, diagnosis,
            classification, failure_modes, str(ai_maturity), results
        )
        
        # Cache results
        self.cache.set(inputs, results)
        
        # Initialize interaction state
        self.interaction_state.update_state(company_name, 'HOOK_SENT')
        
        logger.info("Production pipeline completed")
        return results
    
    def _reconstruct_pipeline_results(self, company_data: Dict[str, Any]):
        """Reconstruct PipelineResults from stored company data without recomputation"""
        signals = company_data['extracted_signals']
        
        # Reconstruct patterns if available
        patterns = None
        if 'patterns' in company_data and company_data['patterns']:
            patterns = company_data['patterns']
        else:
            # Generate patterns for backward compatibility
            patterns = self.pattern_engine.detect_patterns(signals, company_data['company_name'])
        
        return PipelineResults(
            signals=signals,
            diagnosis=company_data['diagnostic_summary'],
            hook=company_data.get('hook', 'Hook not stored in legacy record'),
            audit=company_data.get('audit', 'Audit not stored in legacy record'), 
            close=company_data.get('close', 'Close not stored in legacy record'),
            patterns=patterns
        )
    
    def _extract_failure_modes(self, diagnosis: str) -> List[str]:
        """Extract failure modes from diagnosis text"""
        failure_modes = []
        lines = diagnosis.split('\n')
        in_failure_section = False
        
        for line in lines:
            if 'Scaling Failure Points:' in line or 'Hidden Bottlenecks:' in line:
                in_failure_section = True
                continue
            elif line.startswith('**') and in_failure_section:
                in_failure_section = False
            elif in_failure_section and line.strip().startswith('-'):
                failure_modes.append(line.strip()[1:].strip())
        
        return failure_modes
    
    def _extract_classification(self, diagnosis: str) -> str:
        """Extract AI classification from diagnosis"""
        classifications = ['AI-Washed', 'AI-Assisted', 'AI-Native', 'Non-AI']
        
        for classification in classifications:
            if classification in diagnosis:
                return classification
        
        return 'Unknown'

def main():
    st.set_page_config(
        page_title="IDPETECH · BA Assistant (Production Advisory System)",
        page_icon="🤖",
        layout="wide"
    )
    
    st.title("🤖 IDPETECH · BA Assistant")
    st.subheader("Production Advisory System with Memory & Pattern Detection")
    
    # Initialize session state
    if 'results' not in st.session_state:
        st.session_state.results = None
    if 'prompts' not in st.session_state:
        ba_temp = BAAssistant("temp")
        st.session_state.prompts = ba_temp._load_prompts()
    if 'ba_assistant' not in st.session_state:
        api_key = st.secrets.get("OPENAI_API_KEY") or st.sidebar.text_input("OpenAI API Key", type="password")
        if api_key:
            st.session_state.ba_assistant = BAAssistant(api_key, st.session_state.prompts)
    
    # Layout: Left panel for inputs, Right panel for outputs
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.header("📋 Company Inputs")
        
        # Load input data if a company was loaded from memory
        default_linkedin = st.session_state.get('loaded_linkedin_url', '')
        default_website = st.session_state.get('loaded_website', '')
        default_job_posting = st.session_state.get('loaded_job_posting', '')
        
        linkedin_url = st.text_area(
            "LinkedIn URL or Company Description",
            value=default_linkedin,
            placeholder="Enter LinkedIn company URL or describe the company...",
            height=100
        )
        
        website = st.text_area(
            "Website or Company Summary",
            value=default_website,
            placeholder="Enter website URL or company summary...",
            height=100
        )
        
        job_posting = st.text_area(
            "Job Posting Text",
            value=default_job_posting,
            placeholder="Paste the job posting content...",
            height=150
        )
        
        if st.button("🚀 Generate AI Readiness Analysis", type="primary"):
            if not all([linkedin_url, website, job_posting]):
                st.error("Please fill in all three input fields.")
            elif 'ba_assistant' not in st.session_state:
                st.error("Please provide OpenAI API key.")
            else:
                website_content = None

                # Scrape the website URL if one was provided
                if _is_url(website):
                    with st.spinner(f"🎯 Comprehensive scraping {website} with JavaScript rendering (FREE)…"):
                        ok, result = scrape_website(website)
                    if ok:
                        website_content = result["main_content"]
                        pages_data = result["pages_scraped"]
                        total_chars = result["total_chars"]
                        page_count = result["page_count"]
                        
                        st.success(f"✅ Website comprehensively scraped with Playwright — {page_count} pages, {total_chars:,} total chars (JavaScript rendered)")
                        
                        # Show comprehensive scraped content in expandable evidence panel
                        with st.expander(f"🔍 **Multi-Page Scraping Evidence** - {page_count} pages scraped", expanded=False):
                            st.markdown("**Source URL:** " + website)
                            st.markdown("**Extraction Method:** Playwright with JavaScript rendering + Intelligent page discovery")
                            st.markdown(f"**Pages Scraped:** {page_count} pages")
                            st.markdown(f"**Total Content:** {total_chars:,} characters")
                            st.markdown("---")
                            
                            # Show overview of scraped pages
                            st.markdown("### 📋 Pages Scraped Overview:")
                            for i, page_data in enumerate(pages_data):
                                st.markdown(f"""
**{i+1}. {page_data['page_type']}** - *{page_data['title']}*
- URL: `{page_data['url']}`
- Content: {page_data['char_count']:,} chars
                                """)
                            
                            st.markdown("---")
                            
                            # Display combined content
                            st.markdown("### 📄 Combined Website Content:")
                            st.markdown("```markdown")
                            st.text(website_content)
                            st.markdown("```")
                            
                            # Show individual page content in tabs
                            st.markdown("---")
                            st.markdown("### 📑 Individual Page Content:")
                            
                            if len(pages_data) > 1:
                                page_tabs = st.tabs([f"{p['page_type']}" for p in pages_data])
                                for i, page_data in enumerate(pages_data):
                                    with page_tabs[i]:
                                        st.markdown(f"**Title:** {page_data['title']}")
                                        st.markdown(f"**URL:** {page_data['url']}")
                                        st.markdown(f"**Content Length:** {page_data['char_count']:,} chars")
                                        st.markdown("---")
                                        st.text_area(
                                            f"Content from {page_data['page_type']}",
                                            value=page_data['content'],
                                            height=300,
                                            key=f"page_content_{i}"
                                        )
                            
                            # Additional metadata
                            st.markdown("---")
                            st.caption("💡 This comprehensive content includes the main page plus automatically discovered relevant pages (About, Team, Products, etc.) that provide full context for AI readiness analysis.")
                    else:
                        error_msg = result.get("error", "Unknown error") if isinstance(result, dict) else result
                        st.warning(f"❌ Playwright multi-page scraping failed: {error_msg}. Proceeding without page content.")

                with st.spinner("Running production advisory pipeline..."):
                    try:
                        inputs = CompanyInputs(linkedin_url, website, job_posting,
                                               website_content=website_content)
                        
                        # Clear any previously loaded company state for fresh analysis
                        if 'loaded_company' in st.session_state:
                            del st.session_state.loaded_company
                        if 'loaded_linkedin_url' in st.session_state:
                            del st.session_state.loaded_linkedin_url
                        if 'loaded_website' in st.session_state:
                            del st.session_state.loaded_website  
                        if 'loaded_job_posting' in st.session_state:
                            del st.session_state.loaded_job_posting
                        
                        st.session_state.results = st.session_state.ba_assistant.run_full_pipeline(inputs)
                        st.success("Analysis complete with pattern detection!")
                    except Exception as e:
                        st.error(f"Error: {str(e)}")
        
        st.divider()
        
        # Company Memory Section
        if 'ba_assistant' in st.session_state:
            with st.expander("📊 Company Memory", expanded=False):
                st.write("**Recent Companies:**")
                companies = st.session_state.ba_assistant.memory.list_companies(10)
                
                if companies:
                    for company in companies:
                        with st.container():
                            col_name, col_class, col_state, col_load = st.columns([2, 1, 1, 1])
                            with col_name:
                                st.write(f"**{company['company_name']}**")
                            with col_class:
                                st.write(f"{company['system_classification']}")
                            with col_state:
                                state = st.session_state.ba_assistant.interaction_state.get_state(company['company_name'])
                                st.write(f"{state or 'NEW'}")
                            with col_load:
                                if st.button("📂 Load", key=f"load_{company['company_name']}_{company.get('timestamp', 'default')}", help="Load this company's analysis"):
                                    # Load full company data without recomputation
                                    full_company_data = st.session_state.ba_assistant.memory.get_company(company['company_name'])
                                    if full_company_data:
                                        # Reconstruct PipelineResults from stored data
                                        st.session_state.results = st.session_state.ba_assistant._reconstruct_pipeline_results(full_company_data)
                                        st.session_state.loaded_company = company['company_name']
                                        
                                        # Load the input data into session state for form population
                                        raw_inputs = full_company_data['raw_inputs']
                                        st.session_state.loaded_linkedin_url = raw_inputs['linkedin_url']
                                        st.session_state.loaded_website = raw_inputs['website'] 
                                        st.session_state.loaded_job_posting = raw_inputs['job_posting']
                                        
                                        st.success(f"Loaded {company['company_name']} with full state")
                                        st.rerun()
                                    else:
                                        st.error("Failed to load company data")
                else:
                    st.write("No companies in memory yet.")
        
        # Prompt Management Section
        with st.expander("⚙️ Prompt Management", expanded=False):
            st.write("**Edit the system prompts used in each stage:**")
            
            prompt_tabs = st.tabs(["Extract", "Diagnose", "Hook", "Audit", "Close"])
            
            with prompt_tabs[0]:
                st.session_state.prompts["extract_signals"] = st.text_area(
                    "Signal Extraction Prompt",
                    value=st.session_state.prompts["extract_signals"],
                    height=200,
                    key="prompt_extract"
                )
            
            with prompt_tabs[1]:
                st.session_state.prompts["diagnose"] = st.text_area(
                    "Diagnostic Prompt",
                    value=st.session_state.prompts["diagnose"],
                    height=200,
                    key="prompt_diagnose"
                )
            
            with prompt_tabs[2]:
                st.session_state.prompts["generate_hook"] = st.text_area(
                    "Hook Generation Prompt",
                    value=st.session_state.prompts["generate_hook"],
                    height=200,
                    key="prompt_hook"
                )
            
            with prompt_tabs[3]:
                st.session_state.prompts["generate_audit"] = st.text_area(
                    "Audit Generation Prompt",
                    value=st.session_state.prompts["generate_audit"],
                    height=200,
                    key="prompt_audit"
                )
            
            with prompt_tabs[4]:
                st.session_state.prompts["generate_close"] = st.text_area(
                    "Close Generation Prompt",
                    value=st.session_state.prompts["generate_close"],
                    height=200,
                    key="prompt_close"
                )
            
            col_reset, col_update, col_save = st.columns(3)
            
            with col_reset:
                if st.button("🔄 Factory Reset"):
                    if 'ba_assistant' in st.session_state:
                        st.session_state.prompts = st.session_state.ba_assistant.reset_to_factory_defaults()
                        st.rerun()
            
            with col_update:
                if st.button("💾 Update Assistant"):
                    if 'ba_assistant' in st.session_state:
                        api_key = st.secrets.get("OPENAI_API_KEY")
                        st.session_state.ba_assistant = BAAssistant(api_key, st.session_state.prompts)
                        st.success("Assistant updated with new prompts!")
            
            with col_save:
                if st.button("⭐ Save as Defaults"):
                    if 'ba_assistant' in st.session_state:
                        if st.session_state.ba_assistant.save_custom_prompts(st.session_state.prompts):
                            st.success("Prompts saved as new defaults!")
                        else:
                            st.error("Failed to save prompts")
            
            st.info("💡 **Tip:** Edit prompts above, click 'Update Assistant', then 'Save as Defaults' to make your changes permanent for future sessions.")
    
    with col2:
        if st.session_state.results:
            results = st.session_state.results
            # Use loaded company name if available, otherwise get from signals
            if 'loaded_company' in st.session_state and st.session_state.loaded_company:
                company_name = st.session_state.loaded_company
            else:
                company_name = results.signals.get('company_name', 'Unknown Company')
            
            # Display header with company name
            st.header(f"📊 Analysis Results - {company_name}")
            
            pdf_generator = PDFGenerator()
            word_generator = WordGenerator()
            timestamp = datetime.now().strftime('%Y%m%d_%H%M')
            
            # Helper function for individual downloads
            def create_download_buttons(section_title: str, content: str, file_prefix: str):
                col_md, col_pdf, col_word = st.columns(3)
                
                with col_md:
                    markdown_content = f"""# {section_title}
Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}
Company: {company_name}

{content}
"""
                    st.download_button(
                        label="📄 MD",
                        data=markdown_content,
                        file_name=f"{file_prefix}_{timestamp}.md",
                        mime="text/markdown",
                        key=f"md_{file_prefix}"
                    )
                
                with col_pdf:
                    try:
                        pdf_bytes = pdf_generator.generate_section_pdf(section_title, content, company_name)
                        st.download_button(
                            label="📄 PDF",
                            data=pdf_bytes,
                            file_name=f"{file_prefix}_{timestamp}.pdf",
                            mime="application/pdf",
                            key=f"pdf_{file_prefix}"
                        )
                    except Exception as e:
                        st.error(f"PDF generation failed: {str(e)}")
                
                with col_word:
                    try:
                        word_bytes = word_generator.generate_section_word(section_title, content, company_name)
                        st.download_button(
                            label="📄 DOCX",
                            data=word_bytes,
                            file_name=f"{file_prefix}_{timestamp}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"word_{file_prefix}"
                        )
                    except Exception as e:
                        st.error(f"Word generation failed: {str(e)}")
            
            # Pattern Insights Section (NEW)
            if results.patterns and results.patterns.get('similar_companies'):
                with st.expander("🔍 Cross-Company Pattern Analysis", expanded=True):
                    patterns = results.patterns
                    
                    st.write("**Similar Companies Found:**")
                    for similar in patterns['similar_companies']:
                        st.write(f"• **{similar['company_name']}** (Score: {similar['similarity_score']}) - {similar['classification']}")
                        if similar['tech_overlap']:
                            st.write(f"  Tech overlap: {', '.join(similar['tech_overlap'])}")
                    
                    if patterns['recurring_failure_modes']:
                        st.write("**Recurring Failure Modes:**")
                        for failure in patterns['recurring_failure_modes']:
                            st.write(f"• {failure['failure_mode']} (seen {failure['frequency']}x)")
                    
                    st.write("**Pattern Summary:**")
                    st.write(patterns['common_patterns'])
            
            # Hook Section
            st.subheader("🎯 First Engagement Hook")
            st.info(results.hook)
            create_download_buttons("First Engagement Hook", results.hook, "hook")
            
            st.divider()
            
            # Diagnostic Section  
            st.subheader("🔍 Diagnostic Insight")
            st.markdown(results.diagnosis)
            create_download_buttons("Diagnostic Insight", results.diagnosis, "diagnosis")
            
            st.divider()
            
            # Audit Section
            st.subheader("📋 AI Readiness Audit")
            st.markdown(results.audit)
            create_download_buttons("AI Readiness Audit", results.audit, "audit")
            
            st.divider()
            
            # Close Section
            st.subheader("🤝 Conversation Close")
            st.write(results.close)
            create_download_buttons("Conversation Close", results.close, "close")
            
            st.divider()
            
            # Interaction State Management (NEW)
            if 'ba_assistant' in st.session_state:
                with st.expander("📈 Interaction State Management", expanded=False):
                    current_state = st.session_state.ba_assistant.interaction_state.get_state(company_name)
                    st.write(f"**Company:** {company_name}")
                    st.write(f"**Current State:** {current_state or 'NEW'}")
                    
                    # Show state indicator with color
                    if current_state:
                        if current_state == 'HOOK_SENT':
                            st.info(f"📤 {current_state}")
                        elif current_state == 'ENGAGED':
                            st.success(f"✅ {current_state}")
                        elif current_state in ['QUALIFIED', 'AUDIT_SHARED']:
                            st.success(f"🎯 {current_state}")
                        elif current_state == 'DROPPED':
                            st.error(f"❌ {current_state}")
                        else:
                            st.warning(f"⏳ {current_state}")
                    
                    st.write("**Update State:**")
                    new_state = st.selectbox(
                        "Select new interaction state:",
                        st.session_state.ba_assistant.interaction_state.STATES,
                        index=st.session_state.ba_assistant.interaction_state.STATES.index(current_state) if current_state in st.session_state.ba_assistant.interaction_state.STATES else 0,
                        key="state_selector"
                    )
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("Update State", key="update_state_btn"):
                            if st.session_state.ba_assistant.interaction_state.update_state(company_name, new_state):
                                st.success(f"State updated to: {new_state}")
                                st.rerun()
                            else:
                                st.error("Failed to update state")
                    with col2:
                        if current_state and st.button("View State History", key="state_history_btn"):
                            st.info("State history feature coming soon...")
            
            # Complete Report Download Options
            st.subheader("💾 Complete Report")
            st.caption("Download all sections as a single file")
            
            col_md, col_pdf, col_word = st.columns(3)
            
            with col_md:
                complete_markdown = f"""# AI Readiness Audit Report
Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}
Company: {company_name}

## First Engagement Hook
{results.hook}

## Diagnostic Insight  
{results.diagnosis}

## AI Readiness Audit
{results.audit}

## Conversation Close
{results.close}
"""
                st.download_button(
                    label="📄 Complete Report (MD)",
                    data=complete_markdown,
                    file_name=f"complete_ai_readiness_report_{timestamp}.md",
                    mime="text/markdown",
                    key="complete_md"
                )
            
            with col_pdf:
                try:
                    pdf_bytes = pdf_generator.generate_pdf(results, company_name)
                    st.download_button(
                        label="📄 Complete Report (PDF)",
                        data=pdf_bytes,
                        file_name=f"complete_ai_readiness_report_{timestamp}.pdf",
                        mime="application/pdf",
                        key="complete_pdf"
                    )
                except Exception as e:
                    st.error(f"Complete PDF generation failed: {str(e)}")
            
            with col_word:
                try:
                    word_bytes = word_generator.generate_word(results, company_name)
                    st.download_button(
                        label="📄 Complete Report (DOCX)",
                        data=word_bytes,
                        file_name=f"complete_ai_readiness_report_{timestamp}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="complete_word"
                    )
                except Exception as e:
                    st.error(f"Complete Word generation failed: {str(e)}")
        else:
            st.info("👈 Enter company information and click 'Generate' to see the AI readiness analysis with pattern detection.")

if __name__ == "__main__":
    main()